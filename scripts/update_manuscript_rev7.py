#!/usr/bin/env python3
"""Create gMolAI manuscript revision 7 from the authoritative revision 6.

Revision 7 integrates the completed locked-test, MoleculeNet/HIV, and
controlled encoding-speed comparisons.  The locked-test comparison and exact
parameter table already present in revision 6 are retained, while the embedded
radar image and its caption are removed.  The new endpoint and speed results
are added as ordinary Word tables; revision 7 contains no embedded image.

The builder is document-only.  It hash-checks every scientific input, does not
load a molecular encoder, and never modifies revision 6 or a benchmark file.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import tempfile
from collections import Counter
from pathlib import Path
from zipfile import ZipFile

from lxml import etree

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.text.paragraph import Paragraph
except ImportError as error:  # pragma: no cover - build dependency
    raise SystemExit("python-docx==1.2.0 is required to build manuscript rev7") from error


AUTHORITATIVE_REV6_SHA256 = (
    "861adfa58a21707c79ca85bac76a9d78866f341ee44a275f9224cb0d6d416027"
)
TEST_SUMMARY_SHA256 = (
    "e5946cf9b577e9906c810ad810eb3af59cb76a0ac0dea274bb77acb5d3b2812c"
)
MOLECULENET_PRIMARY_SHA256 = (
    "bac6ec56f933d94ef54089130034d065c40a033b3a899590a888011a165d4058"
)
MOLECULENET_COVERAGE_SHA256 = (
    "41cd0505efeff6dec8f301e702f1c05254e219a05437aae382b23ae19aeef793"
)
MOLECULENET_COMPLETE_SHA256 = (
    "f1d716a1e44f150339a96617e529d43c79f65a87b4fb530e37abe605545b5c8a"
)
SPEED_RESULTS_SHA256 = (
    "831fc4e32309bc6822091f2f35f585e4f41d4270d36424af929e3c06e37312b4"
)
SPEED_PROTOCOL_SHA256 = (
    "30d514f9b48534ccd1f35738a38fe8fadecea5d6beb391a4143859a50c8953bc"
)
SPEED_COMPLETE_SHA256 = (
    "219cd4886c46793ba9c7baccd59501ba71e80089d1fd365c624d1bbce043db21"
)

MODEL_ORDER = (
    "gmolai",
    "morgan",
    "molai",
    "molformer",
    "smi_ted",
    "molclr_gin",
    "kermt_v2",
)
MODEL_LABELS = {
    "gmolai": "gMolAI",
    "morgan": "Morgan radius-2",
    "molai": "MolAI epoch 6",
    "molformer": "MoLFormer",
    "smi_ted": "SMI-TED-Light",
    "molclr_gin": "MolCLR-GIN",
    "kermt_v2": "KERMT v2",
}
ENDPOINT_HEADERS = {
    "gmolai": "gMolAI",
    "morgan": "Morgan",
    "molai": "MolAI",
    "molformer": "MoLFormer",
    "smi_ted": "SMI-TED",
    "molclr_gin": "MolCLR",
    "kermt_v2": "KERMT v2",
}
DATASET_ORDER = (
    "bace",
    "bbbp",
    "esol",
    "freesolv",
    "lipophilicity",
    "hiv",
)
DATASET_LABELS = {
    "bace": "BACE AUC",
    "bbbp": "BBBP AUC",
    "esol": "ESOL RMSE",
    "freesolv": "FreeSolv RMSE",
    "lipophilicity": "Lipo. RMSE",
    "hiv": "HIV AUC",
}
COMMON_COUNTS = {
    "bace": 1_502,
    "bbbp": 1_833,
    "esol": 1_116,
    "freesolv": 638,
    "lipophilicity": 4_187,
    "hiv": 36_228,
}
GMOLAI_RANKS = {
    "bace": 5,
    "bbbp": 4,
    "esol": 1,
    "freesolv": 1,
    "lipophilicity": 2,
    "hiv": 3,
}
BATCH_SIZES = (64, 128, 256, 512)

REV6_TEST_PARAGRAPH = (
    "The post-selection frozen-encoder comparison used 9,958 common "
    "probe-training molecules and 49,844 common molecules from the locked "
    "internal test panel. Fig. 3 presents six complementary metrics, each "
    "divided by the highest observed value on that spoke; this visual "
    "normalization defines neither an aggregate score nor a significance "
    "test. gMolAI had the highest topology mean R² (0.9705) and "
    "scaffold-disjoint topology mean R² (0.9729). Morgan had the strongest "
    "recurring-scaffold clustering (ARI 0.3956; NMI 0.7928), while KERMT v2 "
    "was the closest learned comparator for clustering (ARI 0.3620; NMI "
    "0.7548). Thus, the matched comparison supports complementary strengths "
    "rather than universal superiority."
)
REV6_FIGURE_CAPTION = (
    "Figure 3 | Frozen-encoder comparison on the all-model common locked "
    "internal-test panel. Six raw metrics were independently divided by their "
    "maximum across the seven representations so that the best observed value "
    "on each spoke equals 1.0. This best-relative display does not define a "
    "composite score. Topology probes used 9,958 common training-partition "
    "molecules and 49,844 common locked-test molecules; clustering used 16,360 "
    "recurring-scaffold molecules. Exact raw values are retained in the source "
    "artifact."
)
REV7_TEST_PARAGRAPH = (
    "The post-selection frozen-encoder comparison used 9,958 common "
    "probe-training molecules and 49,844 common molecules from the locked "
    "internal test panel. Six complementary metrics were evaluated separately; "
    "no aggregate score or significance test was defined. gMolAI had the "
    "highest topology mean R² (0.9705) and scaffold-disjoint topology mean R² "
    "(0.9729). Morgan had the strongest recurring-scaffold clustering (ARI "
    "0.3956; NMI 0.7928), while KERMT v2 was the closest learned comparator for "
    "clustering (ARI 0.3620; NMI 0.7548). Thus, the matched comparison supports "
    "complementary strengths rather than universal superiority."
)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def require_hash(path: Path, expected: str, label: str) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)
    observed = sha256_file(path)
    if observed != expected:
        raise RuntimeError(f"{label} hash changed: {observed}; expected {expected}")


def read_csv(path: Path, expected: str, label: str) -> list[dict[str, str]]:
    require_hash(path, expected, label)
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def read_json(path: Path, expected: str, label: str) -> dict:
    require_hash(path, expected, label)
    with path.open(encoding="utf-8") as handle:
        return json.load(handle)


def validate_test_summary(path: Path) -> None:
    rows = read_csv(path, TEST_SUMMARY_SHA256, "locked-test summary")
    if tuple(row["model"] for row in rows) != MODEL_ORDER:
        raise RuntimeError("Locked-test model order changed")
    by_model = {row["model"]: row for row in rows}
    expected = {
        ("gmolai", "topology_mean_r2"): 0.9705115056488679,
        ("gmolai", "scaffold_disjoint_mean_r2"): 0.9728505497473741,
        ("morgan", "clustering_ari"): 0.3955703042792596,
        ("morgan", "clustering_nmi"): 0.7927995386560196,
        ("kermt_v2", "clustering_ari"): 0.3619956070084435,
        ("kermt_v2", "clustering_nmi"): 0.754795257961806,
    }
    for (model, field), value in expected.items():
        if float(by_model[model][field]) != value:
            raise RuntimeError(f"Locked-test {model} {field} changed")
    if {int(row["train_rows"]) for row in rows} != {9_958}:
        raise RuntimeError("Locked-test common training count changed")
    if {int(row["test_rows"]) for row in rows} != {49_844}:
        raise RuntimeError("Locked-test common test count changed")
    if {int(row["clustering_rows"]) for row in rows} != {16_360}:
        raise RuntimeError("Locked-test clustering count changed")


def load_endpoint_results(
    primary_path: Path,
    coverage_path: Path,
    complete_path: Path,
) -> dict[tuple[str, str], dict[str, str]]:
    rows = read_csv(
        primary_path,
        MOLECULENET_PRIMARY_SHA256,
        "MoleculeNet common-panel primary metrics",
    )
    lookup = {(row["dataset"], row["model"]): row for row in rows}
    if len(rows) != 42 or len(lookup) != 42:
        raise RuntimeError("Expected six datasets by seven endpoint representations")
    for dataset in DATASET_ORDER:
        dataset_rows = [lookup[(dataset, model)] for model in MODEL_ORDER]
        if {int(row["common_molecules"]) for row in dataset_rows} != {
            COMMON_COUNTS[dataset]
        }:
            raise RuntimeError(f"Common molecule count changed for {dataset}")
        if int(lookup[(dataset, "gmolai")]["rank"]) != GMOLAI_RANKS[dataset]:
            raise RuntimeError(f"gMolAI rank changed for {dataset}")
        task = lookup[(dataset, "gmolai")]["task"]
        metric = lookup[(dataset, "gmolai")]["primary_metric"]
        expected_metric = "roc_auc" if task == "classification" else "rmse"
        if metric != expected_metric:
            raise RuntimeError(f"Primary metric changed for {dataset}")

    coverage = read_csv(
        coverage_path,
        MOLECULENET_COVERAGE_SHA256,
        "MoleculeNet coverage",
    )
    all_rows = {row["model"]: row for row in coverage if row["dataset"] == "all"}
    if set(all_rows) != set(MODEL_ORDER):
        raise RuntimeError("MoleculeNet all-dataset coverage models changed")
    if {int(row["rows"]) for row in all_rows.values()} != {46_551}:
        raise RuntimeError("MoleculeNet prepared-row total changed")
    expected_accepted = {
        "gmolai": 46_551,
        "morgan": 46_551,
        "molai": 45_505,
        "molformer": 46_488,
        "smi_ted": 46_488,
        "molclr_gin": 46_551,
        "kermt_v2": 46_551,
    }
    for model, accepted in expected_accepted.items():
        if int(all_rows[model]["accepted"]) != accepted:
            raise RuntimeError(f"MoleculeNet coverage changed for {model}")

    complete = read_json(
        complete_path,
        MOLECULENET_COMPLETE_SHA256,
        "MoleculeNet completion seal",
    )
    if complete.get("status") != "complete" or complete.get("common_rows") != 45_504:
        raise RuntimeError("MoleculeNet completion seal is not the authoritative result")
    if complete.get("models") != list(MODEL_ORDER):
        raise RuntimeError("MoleculeNet completion model order changed")
    return lookup


def load_speed_results(
    results_path: Path,
    protocol_path: Path,
    complete_path: Path,
) -> dict[tuple[str, int], dict[str, str]]:
    rows = read_csv(results_path, SPEED_RESULTS_SHA256, "speed results")
    lookup = {(row["model"], int(row["batch_size"])): row for row in rows}
    if len(rows) != 28 or len(lookup) != 28:
        raise RuntimeError("Expected seven encoders by four speed conditions")
    for model in MODEL_ORDER:
        for batch_size in BATCH_SIZES:
            row = lookup[(model, batch_size)]
            if int(row["rows"]) != 49_844:
                raise RuntimeError("Speed panel count changed")
            if row["host"] != "n54" or row["slurm_job_id"] != "1230738":
                raise RuntimeError("Speed execution identity changed")
            equivalent = row["within_tolerance_of_reference"] == "True"
            if model == "kermt_v2" and batch_size != 64:
                if equivalent:
                    raise RuntimeError("KERMT unexpectedly passed cross-batch equivalence")
            elif not equivalent:
                raise RuntimeError(f"{model} batch {batch_size} failed equivalence")

    protocol = read_json(protocol_path, SPEED_PROTOCOL_SHA256, "speed protocol")
    if tuple(protocol["execution"]["batch_sizes"]) != BATCH_SIZES:
        raise RuntimeError("Speed batch-size protocol changed")
    if protocol["execution"]["cpu_workers"] != 48:
        raise RuntimeError("Qualified gMolAI worker count changed")

    complete = read_json(
        complete_path,
        SPEED_COMPLETE_SHA256,
        "speed completion seal",
    )
    if complete.get("status") != (
        "complete_with_declared_kermt_cross_batch_nonconformance"
    ):
        raise RuntimeError("Speed completion status changed")
    if complete.get("panel_rows") != 49_844:
        raise RuntimeError("Speed completion panel count changed")
    if complete.get("models") != list(MODEL_ORDER):
        raise RuntimeError("Speed completion model order changed")
    return lookup


def find_exact_paragraph(document, text: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact paragraph {text!r}, found {len(matches)}")
    return matches[0]


def find_paragraph(document, prefix: str) -> Paragraph:
    matches = [
        paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}, found {len(matches)}")
    return matches[0]


def paragraph_after(anchor: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def paragraph_before(anchor: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is None:
        raise RuntimeError("Paragraph is already detached")
    parent.remove(element)


def table_matrix(table) -> tuple[tuple[str, ...], ...]:
    return tuple(tuple(cell.text for cell in row.cells) for row in table.rows)


def omml_hashes(document) -> Counter[str]:
    hashes: Counter[str] = Counter()
    for tag in (qn("m:oMath"), qn("m:oMathPara")):
        for node in document.element.body.iter(tag):
            hashes[hashlib.sha256(etree.tostring(node)).hexdigest()] += 1
    return hashes


def old_text_is_subsequence(old: list[str], new: list[str]) -> bool:
    position = 0
    for text in new:
        if position < len(old) and text == old[position]:
            position += 1
    return position == len(old)


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_cant_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def set_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 7.0,
    align=WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    properties = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", 25), ("left", 30), ("bottom", 25), ("right", 30)):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    properties.append(margins)


def insert_table_before(
    document,
    anchor: Paragraph,
    matrix: list[list[str]],
    widths: tuple[float, ...],
    *,
    bold_cells: set[tuple[int, int]] | None = None,
    font_size: float = 7.0,
):
    if not matrix or len(widths) != len(matrix[0]):
        raise ValueError("Table matrix and widths disagree")
    if any(len(row) != len(widths) for row in matrix):
        raise ValueError("Table rows have inconsistent widths")
    table = document.add_table(rows=len(matrix), cols=len(widths))
    table.style = document.tables[0].style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    bold_cells = bold_cells or set()
    for row_index, values in enumerate(matrix):
        row = table.rows[row_index]
        set_cant_split(row)
        if row_index == 0:
            set_repeat_header(row)
        for column_index, (value, width) in enumerate(zip(values, widths)):
            set_cell(
                row.cells[column_index],
                value,
                bold=(row_index == 0 or (row_index, column_index) in bold_cells),
                size=font_size,
            )
            row.cells[column_index].width = Inches(width)
    anchor._p.addprevious(table._tbl)
    return table


def endpoint_table_matrix(
    lookup: dict[tuple[str, str], dict[str, str]],
) -> tuple[list[list[str]], set[tuple[int, int]]]:
    matrix = [["Dataset / metric", *(ENDPOINT_HEADERS[m] for m in MODEL_ORDER)]]
    bold_cells: set[tuple[int, int]] = set()
    for row_index, dataset in enumerate(DATASET_ORDER, start=1):
        values = [DATASET_LABELS[dataset]]
        for column_index, model in enumerate(MODEL_ORDER, start=1):
            row = lookup[(dataset, model)]
            values.append(
                f"{float(row['mean']):.4f} ± {float(row['population_std']):.4f}"
            )
            if int(row["rank"]) == 1:
                bold_cells.add((row_index, column_index))
        matrix.append(values)
    return matrix, bold_cells


def speed_table_matrix(
    lookup: dict[tuple[str, int], dict[str, str]],
) -> tuple[list[list[str]], set[tuple[int, int]]]:
    matrix = [[
        "Encoder",
        "Device",
        "Batch 64",
        "Batch 128",
        "Batch 256",
        "Batch 512",
        "Equivalent",
    ]]
    bold_cells: set[tuple[int, int]] = set()
    for row_index, model in enumerate(MODEL_ORDER, start=1):
        rows = [lookup[(model, batch_size)] for batch_size in BATCH_SIZES]
        equivalent = "No" if model == "kermt_v2" else "Yes"
        matrix.append([
            MODEL_LABELS[model] + ("†" if model == "kermt_v2" else ""),
            rows[0]["device_class"].upper(),
            *(f"{float(row['rows_per_second']):,.2f}" for row in rows),
            equivalent,
        ])
        if model == "gmolai":
            bold_cells.update((row_index, column) for column in range(2, 6))
    return matrix, bold_cells


def assert_archive_has_no_images(path: Path) -> None:
    with ZipFile(path) as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            raise RuntimeError(f"Corrupt DOCX member: {bad_member}")
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        if media:
            raise RuntimeError(f"rev7 unexpectedly contains media: {media}")
        prohibited = (
            b"<w:drawing>",
            b"<w:drawing ",
            b"<w:pict>",
            b"<w:pict ",
            b"<a:blip>",
            b"<a:blip ",
            b"<v:imagedata>",
            b"<v:imagedata ",
        )
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith((".xml", ".rels")):
                continue
            payload = archive.read(name)
            if any(marker in payload for marker in prohibited):
                raise RuntimeError(f"rev7 contains an image element in {name}")
            if b"relationships/image" in payload:
                raise RuntimeError(f"rev7 contains an image relationship in {name}")


def validate_serialized(
    path: Path,
    old_tables: list[tuple[tuple[str, ...], ...]],
    old_omml: Counter[str],
    old_paragraphs: list[str],
    endpoint_table: tuple[tuple[str, ...], ...],
    speed_table: tuple[tuple[str, ...], ...],
) -> None:
    reopened = Document(path)
    if len(reopened.tables) != 15:
        raise RuntimeError(f"rev7 must contain 15 tables, found {len(reopened.tables)}")
    if len(reopened.inline_shapes) != 0:
        raise RuntimeError("rev7 must contain no inline image")
    if [table_matrix(table) for table in reopened.tables[:13]] != old_tables:
        raise RuntimeError("Serialized rev7 changed an existing rev6 table")
    if table_matrix(reopened.tables[13]) != endpoint_table:
        raise RuntimeError("Serialized rev7 endpoint table changed")
    if table_matrix(reopened.tables[14]) != speed_table:
        raise RuntimeError("Serialized rev7 speed table changed")
    if omml_hashes(reopened) != old_omml:
        raise RuntimeError("Serialized rev7 changed an existing equation")
    new_paragraphs = [p.text for p in reopened.paragraphs if p.text]
    if not old_text_is_subsequence(old_paragraphs, new_paragraphs):
        raise RuntimeError("Serialized rev7 changed unedited rev6 paragraph text or order")
    captions = [
        int(match.group(1))
        for paragraph in reopened.paragraphs
        if (match := re.match(r"^Table (\d+) \|", paragraph.text))
    ]
    if captions != list(range(1, 16)):
        raise RuntimeError(f"Non-sequential rev7 table captions: {captions}")
    full_text = "\n".join(new_paragraphs)
    if "Figure 3" in full_text or "Fig. 3" in full_text:
        raise RuntimeError("rev7 retained the removed embedded Figure 3 reference")
    required = (
        "1.6.5. Frozen seven-representation common-panel comparison",
        "1.9.1. Controlled single-GPU encoding-throughput benchmark",
        "2.1.9. Frozen seven-representation MoleculeNet and HIV comparison",
        "2.1.10. Controlled single-GPU encoding throughput",
        "45,504-molecule all-model intersection",
        "58,330.38 molecules per second",
        "5.36-fold the Morgan throughput",
        "10.04-fold the fastest other representation-equivalent neural encoder",
        "minimum cosine similarity of 0.982395",
        "Table 14",
        "Table 15",
    )
    missing = [value for value in required if value not in full_text]
    if missing:
        raise RuntimeError(f"Required rev7 content is missing: {missing}")
    assert_archive_has_no_images(path)


def build(args: argparse.Namespace) -> None:
    source = Path(args.input).resolve()
    destination = Path(args.output).resolve()
    test_path = Path(args.test_summary).resolve()
    endpoint_path = Path(args.moleculenet_primary).resolve()
    coverage_path = Path(args.moleculenet_coverage).resolve()
    endpoint_complete_path = Path(args.moleculenet_complete).resolve()
    speed_path = Path(args.speed_results).resolve()
    speed_protocol_path = Path(args.speed_protocol).resolve()
    speed_complete_path = Path(args.speed_complete).resolve()
    if source == destination:
        raise ValueError("rev7 output must differ from the retained rev6 input")
    if destination.exists() and not args.overwrite:
        raise FileExistsError(f"Output exists; pass --overwrite to replace it: {destination}")

    require_hash(source, AUTHORITATIVE_REV6_SHA256, "authoritative rev6")
    validate_test_summary(test_path)
    endpoint_lookup = load_endpoint_results(
        endpoint_path, coverage_path, endpoint_complete_path
    )
    speed_lookup = load_speed_results(
        speed_path, speed_protocol_path, speed_complete_path
    )

    document = Document(source)
    if len(document.tables) != 13 or len(document.inline_shapes) != 1:
        raise RuntimeError("Authoritative rev6 must contain 13 tables and one image")
    old_tables = [table_matrix(table) for table in document.tables]
    old_omml = omml_hashes(document)
    excluded_text = {REV6_TEST_PARAGRAPH, REV6_FIGURE_CAPTION}
    old_paragraphs = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text not in excluded_text
    ]

    test_paragraph = find_exact_paragraph(document, REV6_TEST_PARAGRAPH)
    test_paragraph.text = REV7_TEST_PARAGRAPH
    figure_caption = find_exact_paragraph(document, REV6_FIGURE_CAPTION)
    drawing_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.xpath(".//w:drawing") or paragraph._p.xpath(".//w:pict")
    ]
    if len(drawing_paragraphs) != 1:
        raise RuntimeError(f"Expected one rev6 drawing paragraph, found {len(drawing_paragraphs)}")
    remove_paragraph(drawing_paragraphs[0])
    remove_paragraph(figure_caption)
    for relationship_id, relationship in list(document.part.rels.items()):
        if relationship.reltype == RT.IMAGE:
            document.part.drop_rel(relationship_id)

    table5_anchor = find_exact_paragraph(
        document,
        "Table 5 | Frozen promotion criteria used in the complete retrospective checkpoint audit.",
    )
    paragraph_before(
        table5_anchor,
        "1.6.5. Frozen seven-representation common-panel comparison",
        "Heading 3",
    )
    paragraph_before(
        table5_anchor,
        "After promotion, the unchanged endpoint workflow was extended retrospectively to "
        "seven fixed representations: gMolAI, Morgan radius-2 fingerprints, MolAI epoch 6, "
        "MoLFormer, SMI-TED-Light, MolCLR-GIN and KERMT v2. Their native output widths were "
        "384, 2,048, 512, 768, 768, 512 and 512 dimensions, respectively. No neural encoder "
        "was trained or fine-tuned, no representation was selected on these results, and no "
        "PCA, cross-model projection or endpoint-specific feature selection was applied.",
        "Body Text",
    )
    paragraph_before(
        table5_anchor,
        "Every adapter screened the complete prepared panel before fitting a downstream "
        "probe. The primary paired analysis used the 45,504-molecule all-model intersection: "
        "1,502 BACE, "
        "1,833 BBBP, 1,116 ESOL, 638 FreeSolv, 4,187 Lipophilicity and 36,228 HIV molecules, "
        "for 45,504 molecular occurrences in total. Existing outer and inner scaffold roles "
        "were retained after intersection; no split was regenerated. Each representation "
        "therefore used the same fold-local scaling, hyperparameter candidates, inherited "
        "split identities and linear-probe family within an endpoint.",
        "Body Text",
    )

    software_anchor = find_exact_paragraph(
        document, "1.10. Software environment and computational reproducibility"
    )
    paragraph_before(
        software_anchor,
        "1.9.1. Controlled single-GPU encoding-throughput benchmark",
        "Heading 3",
    )
    paragraph_before(
        software_anchor,
        "The seven encoders were measured sequentially in one Slurm job on Arrhenius node "
        "n54 using the same ordered 49,844-molecule all-model common locked-test panel. One "
        "NVIDIA GH200 120GB GPU was visible to each neural encoder, while Morgan remained a "
        "CPU implementation in the same 48-CPU allocation. Batch sizes 64, 128, 256 and 512 "
        "were evaluated in fixed order after one untimed warm-up batch. The optimized gMolAI "
        "pipeline used its qualified 48-worker RDKit pool; every competitor retained its "
        "native preprocessing under the same CPU ceiling.",
        "Body Text",
    )
    paragraph_before(
        software_anchor,
        "The timer began with canonical SMILES resident in memory and ended when every "
        "ordered FP32 vector was resident in host memory. It included model-specific parsing "
        "or graph construction, host-to-device transfer, frozen forward inference, "
        "device-to-host transfer and output-matrix materialization. It excluded SIF and model "
        "loading, input-file reading, worker startup, warm-up, validation, hashing and disk "
        "serialization. Batch-128/256/512 outputs were checked against batch 64 at minimum "
        "cosine similarity 0.9999 and maximum per-row relative-L2 delta 0.005. Each condition "
        "had one measured pass, so the benchmark provides descriptive point measurements "
        "without confidence intervals.",
        "Body Text",
    )

    discussion_anchor = find_exact_paragraph(document, "2.2. Discussion")
    paragraph_before(
        discussion_anchor,
        "2.1.9. Frozen seven-representation MoleculeNet and HIV comparison",
        "Heading 3",
    )
    paragraph_before(
        discussion_anchor,
        "Across the six endpoints, 46,551 prepared molecular occurrences were screened and "
        "45,504 (97.751%) entered the all-model intersection. gMolAI, Morgan, MolCLR-GIN and "
        "KERMT v2 accepted every prepared row. MolAI accepted 45,505 rows, while MoLFormer "
        "and SMI-TED-Light each accepted 46,488. Unsupported characters, released token "
        "limits and lossless-tokenization checks were enforced without silent truncation or "
        "unknown-token substitution. Dataset-specific common counts are reported with the "
        "paired results in Table 14.",
        "Normal",
    )
    paragraph_before(
        discussion_anchor,
        "The all-model comparison sharpened the endpoint-dependent pattern. gMolAI ranked "
        "first by mean on ESOL and FreeSolv, second on Lipophilicity, third on confirmatory "
        "HIV, fourth on BBBP and fifth on BACE. Its clearest favorable result was ESOL, where "
        "its mean RMSE of 0.7314 beat Morgan, MoLFormer, MolCLR-GIN and KERMT v2 on all ten "
        "paired splits, and MolAI and SMI-TED-Light on 9/10. On FreeSolv, gMolAI "
        "(1.3055) and SMI-TED-Light (1.3077) were practically tied; on Lipophilicity, "
        "MoLFormer (0.8043) and gMolAI (0.8086) split the paired comparisons 5/5.",
        "Normal",
    )
    caption14 = paragraph_before(
        discussion_anchor,
        "Table 14 | Seven-representation common-panel MoleculeNet and HIV endpoint comparison.",
        "Table Caption",
    )
    caption14.runs[0].bold = True
    caption14.paragraph_format.keep_with_next = True
    endpoint_matrix, endpoint_bold = endpoint_table_matrix(endpoint_lookup)
    endpoint_table = insert_table_before(
        document,
        discussion_anchor,
        endpoint_matrix,
        (0.88, 0.80, 0.80, 0.76, 0.82, 0.80, 0.76, 0.80),
        bold_cells=endpoint_bold,
        font_size=6.4,
    )
    endpoint_note = paragraph_before(
        discussion_anchor,
        "Values are mean ± descriptive population s.d. (ddof=0) over the same ten inherited, "
        "overlapping outer scaffold splits. ROC-AUC is higher-is-better; RMSE is lower-is-"
        "better. Bold marks the best mean only and is not a significance claim. Common "
        "molecule counts were BACE 1,502; BBBP 1,833; ESOL 1,116; FreeSolv 638; "
        "Lipophilicity 4,187; and HIV 36,228. Native representation dimensions differed.",
        "Table Note",
    )
    endpoint_table._tbl.addnext(endpoint_note._p)
    paragraph_before(
        discussion_anchor,
        "On the 36,228-molecule common HIV panel, MoLFormer led mean ROC-AUC at 0.7572, "
        "followed by KERMT v2 at 0.7544 and gMolAI at 0.7507; gMolAI exceeded Morgan "
        "(0.7357) on 8/10 paired splits. These values differ from the preceding full-panel "
        "gMolAI/Morgan HIV result because Table 14 uses the stricter seven-model intersection. "
        "The two analyses answer different coverage questions and were not mixed. Overall, "
        "sequence encoders led BACE and BBBP, whereas gMolAI was strongest or near-strongest "
        "on the descriptor-related regression endpoints. The defensible interpretation is "
        "complementarity, not universal superiority.",
        "Normal",
    )

    paragraph_before(
        discussion_anchor,
        "2.1.10. Controlled single-GPU encoding throughput",
        "Heading 3",
    )
    gmolai_512 = float(speed_lookup[("gmolai", 512)]["rows_per_second"])
    morgan_512 = float(speed_lookup[("morgan", 512)]["rows_per_second"])
    other_neural_512 = max(
        float(speed_lookup[(model, 512)]["rows_per_second"])
        for model in ("molai", "molformer", "smi_ted", "molclr_gin")
    )
    paragraph_before(
        discussion_anchor,
        "The speed-optimized gMolAI implementation was the highest-throughput path at every "
        "tested batch size under the fixed execution envelope (Table 15). Throughput increased "
        f"from {float(speed_lookup[('gmolai', 64)]['rows_per_second']):,.2f} molecules per "
        f"second at batch 64 to {gmolai_512:,.2f} molecules per second at batch 512. At batch "
        f"512 this was {gmolai_512 / morgan_512:.2f}-fold the Morgan throughput and "
        f"{gmolai_512 / other_neural_512:.2f}-fold the fastest other representation-equivalent "
        "neural encoder, MoLFormer. Morgan's value is CPU throughput and must not be described "
        "as GPU-forward performance.",
        "Normal",
    )
    caption15 = paragraph_before(
        discussion_anchor,
        "Table 15 | Controlled end-to-end encoding throughput on the common locked-test panel.",
        "Table Caption",
    )
    caption15.runs[0].bold = True
    caption15.paragraph_format.keep_with_next = True
    speed_matrix, speed_bold = speed_table_matrix(speed_lookup)
    speed_table = insert_table_before(
        document,
        discussion_anchor,
        speed_matrix,
        (1.20, 0.52, 0.84, 0.84, 0.84, 0.84, 0.78),
        bold_cells=speed_bold,
        font_size=6.8,
    )
    speed_note = paragraph_before(
        discussion_anchor,
        "Throughput is molecules s⁻¹ for one complete measured pass after one warm-up batch on "
        "49,844 ordered molecules (Slurm job 1230738, node n54). Neural encoders used one "
        "GH200; Morgan used CPU in the same 48-CPU allocation. The timer included required "
        "preprocessing and transfers but excluded loading, startup, validation and disk output. "
        "There are no confidence intervals. Equivalent denotes cross-batch output equivalence "
        "to the batch-64 representation under the frozen numerical gates.",
        "Table Note",
    )
    speed_table._tbl.addnext(speed_note._p)
    paragraph_before(
        discussion_anchor,
        "† KERMT v2 was exactly repeatable at fixed batch size but failed the unchanged "
        "cross-batch representation-equivalence gate because its native graph builder used a "
        "batch-local adjacency-padding width. Batch 512 versus batch 64 had a minimum cosine "
        "similarity of 0.982395 and a maximum relative-L2 delta of 0.197321. Its speed "
        "points therefore describe computation for batch-dependent native outputs and are not "
        "evidence of scaling for one molecule-only invariant representation. No KERMT source, "
        "checkpoint or prior scientific embedding was changed.",
        "Table Note",
    )

    evidence_anchor = find_paragraph(
        document, "The present results provide three forms of evidence"
    )
    paragraph_after(
        evidence_anchor,
        "The expanded seven-representation endpoint comparison made that complementarity more "
        "specific. gMolAI led the common-panel ESOL result and was effectively tied for the "
        "best FreeSolv mean, while MoLFormer or SMI-TED-Light led BACE, BBBP, Lipophilicity "
        "and HIV. The learned representation therefore adds a strong descriptor-related "
        "regression phenotype without displacing sequence models or Morgan across all tasks.",
        "Normal",
    )

    limitation_anchor = find_paragraph(
        document, "Third, the present MoleculeNet experiments evaluate frozen representations"
    )
    paragraph_after(
        limitation_anchor,
        "Fourth, the controlled speed study contains one measured pass per condition on one "
        "GH200 host and reports implementation-level end-to-end throughput, not intrinsic model "
        "complexity or hardware-independent latency. Its large gMolAI advantage demonstrates "
        "that the released optimized pipeline can utilize this environment efficiently, but "
        "does not supply uncertainty intervals or establish the same ordering on other "
        "architectures. The KERMT batch-dependence finding additionally limits interpretation "
        "of that encoder's throughput across batch sizes.",
        "Normal",
    )

    outlook_anchor = find_paragraph(
        document, "The principal contribution of gMolAI-v2.0 is therefore not a new fingerprint"
    )
    paragraph_after(
        outlook_anchor,
        "For the measured deployment regime, the optimized inference backend materially "
        "strengthens this infrastructure argument: at batch 512 it generated 58,330.38 "
        "molecules per second, 5.36-fold the Morgan throughput and 10.04-fold the fastest other "
        "representation-equivalent neural encoder in the same job. These values are best read "
        "as a concrete capacity result for the pinned GH200 workflow rather than a universal "
        "performance guarantee.",
        "Normal",
    )

    final_anchor = find_paragraph(
        document, "Taken together, the results support a deliberately bounded conclusion"
    )
    paragraph_before(
        final_anchor,
        "The three completed comparator analyses refine the release claim in mutually "
        "consistent ways. The locked internal test panel confirms strong topology accessibility "
        "and exposes Morgan's scaffold-clustering advantage; the common MoleculeNet/HIV panel "
        "shows favorable regression transfer alongside classification endpoints led by other "
        "encoders; and the controlled systems test shows that the qualified optimized gMolAI "
        "implementation can deliver the highest observed throughput under the pinned single-GPU "
        "conditions. None of these findings supports universal scientific or hardware "
        "superiority.",
        "Normal",
    )

    if [table_matrix(table) for table in document.tables[:13]] != old_tables:
        raise RuntimeError("An existing rev6 table changed in memory")
    if omml_hashes(document) != old_omml:
        raise RuntimeError("An existing rev6 equation changed in memory")
    new_paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    if not old_text_is_subsequence(old_paragraphs, new_paragraphs):
        raise RuntimeError("An unedited rev6 paragraph changed or moved in memory")
    if len(document.tables) != 15 or len(document.inline_shapes) != 0:
        raise RuntimeError("rev7 must have 15 tables and no embedded images")

    document.core_properties.subject = (
        "gMolAI-v2.0 manuscript revision 7: completed comparator benchmarks"
    )
    document.core_properties.revision = 7
    destination.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary_name = tempfile.mkstemp(
        prefix=f".{destination.name}.", suffix=".tmp", dir=destination.parent
    )
    os.close(fd)
    temporary = Path(temporary_name)
    try:
        document.save(temporary)
        os.chmod(temporary, source.stat().st_mode & 0o777)
        validate_serialized(
            temporary,
            old_tables,
            old_omml,
            old_paragraphs,
            table_matrix(endpoint_table),
            table_matrix(speed_table),
        )
        if sha256_file(source) != AUTHORITATIVE_REV6_SHA256:
            raise RuntimeError("Authoritative rev6 changed during rev7 construction")
        os.replace(temporary, destination)
    except BaseException:
        try:
            temporary.unlink()
        except FileNotFoundError:
            pass
        raise

    validate_serialized(
        destination,
        old_tables,
        old_omml,
        old_paragraphs,
        table_matrix(endpoint_table),
        table_matrix(speed_table),
    )
    require_hash(source, AUTHORITATIVE_REV6_SHA256, "authoritative rev6 after build")
    validate_test_summary(test_path)
    load_endpoint_results(endpoint_path, coverage_path, endpoint_complete_path)
    load_speed_results(speed_path, speed_protocol_path, speed_complete_path)
    print(
        json.dumps(
            {
                "input": str(source),
                "input_sha256": sha256_file(source),
                "output": str(destination),
                "output_sha256": sha256_file(destination),
                "tables": 15,
                "embedded_images": 0,
                "endpoint_common_rows": 45_504,
                "speed_panel_rows": 49_844,
            },
            indent=2,
            sort_keys=True,
        )
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--test-summary", required=True)
    parser.add_argument("--moleculenet-primary", required=True)
    parser.add_argument("--moleculenet-coverage", required=True)
    parser.add_argument("--moleculenet-complete", required=True)
    parser.add_argument("--speed-results", required=True)
    parser.add_argument("--speed-protocol", required=True)
    parser.add_argument("--speed-complete", required=True)
    parser.add_argument("--overwrite", action="store_true")
    return parser.parse_args()


if __name__ == "__main__":
    build(parse_args())
