#!/usr/bin/env python3
"""Create the editorially reorganized gMolAI manuscript revision 5.

This builder is deliberately document-only.  It reads the authoritative rev4
DOCX, moves existing result paragraphs/tables into evidence-source order, adds
one evidence-roles overview table, and performs fail-closed content checks.  It
does not import or execute any model code.
"""

from __future__ import annotations

import argparse
import hashlib
import os
import re
import tempfile
from collections import Counter
from copy import deepcopy
from pathlib import Path

from lxml import etree

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.text.paragraph import Paragraph
except ImportError as error:  # pragma: no cover - build dependency
    raise SystemExit("python-docx>=1.2 is required to build manuscript rev5") from error


AUTHORITATIVE_REV4_SHA256 = (
    "6a02e6088d29d527160db234fcc76dd9ed0e0e4b180d46e088c08265923d8e37"
)
EXPECTED_REV4_TABLE_CAPTIONS = (
    "Table 1 | Immutable pretraining data contract.",
    "Table 2 | Auxiliary descriptor targets.",
    "Table 3 | Atom and bond feature schema.",
    "Table 4 | Production model and training hyperparameters.",
    "Table 5 | Frozen promotion criteria used in the complete retrospective checkpoint audit.",
    "Table 6 | Complete fail-closed promotion evaluation across retained primary-seed checkpoints.",
    "Table 7 | Pretraining diagnostics continued to improve after the selected checkpoint whereas FreeSolv transfer did not.",
    "Table 8 | Promotion-relevant representation probes for the selected seed-42 checkpoint.",
    "Table 9 | Frozen downstream transfer of the selected 384-D representation, Morgan fingerprints and the 13-descriptor-only control.",
    "Table 10 | Exact checkpoint-resolved downstream-molecule exposure in the seed-42 training stream.",
    "Table 11 | Cross-training-seed replication at step 10,000.",
)
EXPECTED_EXPOSURE = {
    5_000: 28_743_683,
    7_500: 43_109_793,
    10_000: 57_504_265,
    12_500: 71_870_280,
    15_000: 86_236_032,
}
RESULT_HEADINGS = (
    "2.1.1. Pretraining scale, split and actual training exposure",
    "2.1.2. Validation-partition assessment of the selected representation",
    "2.1.3. Complete checkpoint promotion trajectory",
    "2.1.4. Released 384-D representation",
    "2.1.5. External development/promotion benchmarks",
    "2.1.6. Independent seed-43 replication",
    "2.1.7. Locked internal-test evaluation — post-selection confirmation",
    "2.1.8. HIV — independent external post-selection endpoint",
)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def find_paragraph(document, prefix: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning {prefix!r}, found {len(matches)}"
        )
    return matches[0]


def find_exact_paragraph(document, text: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact paragraph {text!r}, found {len(matches)}")
    return matches[0]


def set_text(paragraph: Paragraph, text: str, *, bold: bool = False) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.bold = bold


def replace_in_runs(paragraph: Paragraph, old: str, new: str) -> None:
    matches = [run for run in paragraph.runs if old in run.text]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one run containing {old!r} in {paragraph.text!r}; "
            f"found {len(matches)}"
        )
    matches[0].text = matches[0].text.replace(old, new)


def new_paragraph_before(anchor: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is None:
        raise RuntimeError("Cannot delete a detached paragraph")
    parent.remove(paragraph._p)


def next_table_element(caption: Paragraph):
    element = caption._p.getnext()
    while element is not None and element.tag != qn("w:tbl"):
        if element.tag == qn("w:p") and Paragraph(element, caption._parent).text.strip():
            raise RuntimeError(f"No adjacent table after caption {caption.text!r}")
        element = element.getnext()
    if element is None:
        raise RuntimeError(f"No table after caption {caption.text!r}")
    return element


def table_matrix(table) -> tuple[tuple[str, ...], ...]:
    return tuple(tuple(cell.text for cell in row.cells) for row in table.rows)


def omml_hashes(document) -> Counter[str]:
    hashes: Counter[str] = Counter()
    for node in document.element.body.iter(qn("m:oMath")):
        payload = etree.tostring(node)
        hashes[hashlib.sha256(payload).hexdigest()] += 1
    for node in document.element.body.iter(qn("m:oMathPara")):
        payload = etree.tostring(node)
        hashes[hashlib.sha256(payload).hexdigest()] += 1
    return hashes


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_cant_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    properties.append(marker)


def set_cell(cell, text: str, *, bold: bool = False, size: float = 6.15) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    properties = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", 18), ("left", 20), ("bottom", 18), ("right", 20)):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    properties.append(margins)


def insert_evidence_roles_table(document, anchor: Paragraph) -> tuple[Paragraph, object, Paragraph]:
    heading = new_paragraph_before(
        anchor,
        "1.8. Experimental chronology and data/evidence roles",
        "Heading 2",
    )
    overview = new_paragraph_before(
        anchor,
        "Only the 221,148,895-graph training partition could update encoder weights. The validation "
        "partition could influence checkpoint selection but never model weights. BACE, BBBP, "
        "ESOL, FreeSolv and Lipophilicity formed the selection-conditioned development/promotion "
        "panel. The locked internal test partition was opened only after the seed-42/10,000-step "
        "checkpoint, its train-only calibrator and the 384-D representation definition were "
        "frozen. HIV was a separate external post-selection confirmatory endpoint, not part of "
        "the internal corpus split, and seed 43 was an independent training replication rather "
        "than a model-selection candidate (Table 6).",
        "Body Text",
    )
    caption = new_paragraph_before(
        anchor,
        "Table 6 | Data and evidence roles in training, promotion and post-selection evaluation.",
        "Table Caption",
    )
    set_text(caption, caption.text, bold=True)

    table = document.add_table(rows=1, cols=6)
    table.style = document.tables[4].style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    caption._p.addnext(table._tbl)
    headers = (
        "Evidence role",
        "Source / population",
        "Molecules actually used",
        "Updates neural-network weights?",
        "Can influence seed-42 promotion?",
        "Scientific purpose",
    )
    widths = (0.78, 0.96, 1.58, 0.82, 0.88, 1.38)
    for index, (header, width) in enumerate(zip(headers, widths)):
        set_cell(table.rows[0].cells[index], header, bold=True, size=6.1)
        table.rows[0].cells[index].width = Inches(width)
    set_repeat_header(table.rows[0])
    set_cant_split(table.rows[0])

    rows = (
        (
            "Pretraining training partition",
            "ZINC + PubChem; 221,148,895 graphs",
            "Seed 42 consumed 57,504,265 unique graphs by step 10k; 100,000 fitted its calibrator; 10,000 fitted the topology probe",
            "Yes — consumed training graphs only",
            "Yes — training, calibration and probe fitting",
            "Learn encoder; fit train-only calibration and topology probe",
        ),
        (
            "Pretraining validation partition",
            "943,038 graphs",
            "50,000 main export; 5,000 similarity subset; 11,034 recurring-scaffold clustering subset; separate 50,000 panel (11,011 clustered)",
            "No",
            "Yes — validation can support promotion",
            "Reconstruction, latent, topology, retrieval and clustering evidence",
        ),
        (
            "Development / promotion panel",
            "External BACE, BBBP, ESOL, FreeSolv and Lipophilicity",
            "Post-filter n = 1,513; 1,860; 1,116; 639; 4,198; ten accepted outer scaffold splits each",
            "No",
            "Yes — all five were in the 17-gate decision",
            "Selection-conditioned external transfer evidence",
        ),
        (
            "Locked internal test partition",
            "1,088,766 pretraining graphs",
            "Up to 250,000 for health; separate 50,000 geometry panel; 16,382 recurring-scaffold clustering subset",
            "No",
            "No — opened only after freeze",
            "Protected post-selection internal confirmation and limitation finding",
        ),
        (
            "HIV confirmatory endpoint",
            "External MoleculeNet dataset; post-filter n = 37,225",
            "Ten accepted outer scaffold splits",
            "No",
            "No — evaluated after selection",
            "Independent external post-selection endpoint",
        ),
        (
            "Seed-43 replication",
            "Independent training run on the same pretraining population",
            "Independent run to step 10k; separate calibrator using the same deterministic 100,000-molecule sample; same 50,000-validation and development-panel protocols",
            "Yes — its own encoder only",
            "No — replication, not seed-42 selection",
            "Assess training-seed reproducibility",
        ),
    )
    for values in rows:
        row = table.add_row()
        set_cant_split(row)
        for index, (value, width) in enumerate(zip(values, widths)):
            set_cell(row.cells[index], value)
            row.cells[index].width = Inches(width)

    note = new_paragraph_before(
        anchor,
        "Molecule counts are post-filter counts or explicitly sampled counts as stated. An outer "
        "test fold is only the held-out fold within a MoleculeNet nested scaffold split and is "
        "never synonymous with the locked internal test partition.",
        "Table Note",
    )
    table._tbl.addnext(note._p)
    caption.paragraph_format.keep_with_next = True
    note.paragraph_format.keep_with_next = False
    return heading, table, note


def assert_source_contract(document, source: Path) -> tuple[list[tuple[tuple[str, ...], ...]], Counter[str]]:
    observed_hash = sha256_file(source)
    if observed_hash != AUTHORITATIVE_REV4_SHA256:
        raise RuntimeError(
            f"Refusing non-authoritative rev4 input: {observed_hash}; expected "
            f"{AUTHORITATIVE_REV4_SHA256}"
        )
    captions = tuple(
        p.text for p in document.paragraphs if re.match(r"^Table \d+ \|", p.text)
    )
    if captions != EXPECTED_REV4_TABLE_CAPTIONS:
        raise RuntimeError("rev4 table-caption contract does not match the authoritative source")
    if len(document.tables) != 11:
        raise RuntimeError(f"Expected 11 rev4 tables, found {len(document.tables)}")
    return [table_matrix(t) for t in document.tables], omml_hashes(document)


def build(args: argparse.Namespace) -> None:
    source = Path(args.input).resolve()
    destination = Path(args.output).resolve()
    if source == destination:
        raise ValueError("rev5 output must differ from the retained rev4 input")
    if not source.is_file():
        raise FileNotFoundError(source)

    document = Document(source)
    original_tables, original_omml = assert_source_contract(document, source)

    # Capture every paragraph/table that will be moved before changing captions.
    p = {index: document.paragraphs[index] for index in range(len(document.paragraphs))}
    old_result_tables = {
        number: next_table_element(find_exact_paragraph(document, caption))
        for number, caption in enumerate(EXPECTED_REV4_TABLE_CAPTIONS[5:], start=6)
    }

    # Methods: explicit operational chronology and unambiguous evidence roles.
    set_text(
        p[1],
        "1.1. Study design, experimental chronology and definition of the released representation",
    )
    new_paragraph_before(
        p[5],
        "The operational chronology was: immutable corpus construction and scaffold-aware split; "
        "encoder training; retention of candidate checkpoints; checkpoint-specific calibration "
        "using training-partition molecules; validation-partition assessment; application of the "
        "five-dataset external development/promotion gate; freeze of the seed-42/10,000-step "
        "checkpoint, its calibrator and the 384-D representation; evaluation of the locked internal "
        "test partition; HIV confirmation; and independent seed-43 training replication. Only the "
        "training partition updated neural-network weights.",
        "Body Text",
    )
    set_text(
        p[25],
        "The locked internal test partition was protected: it was not used to fit model "
        "parameters, descriptor normalization, embedding calibration, representation weights, "
        "checkpoint selection or downstream-probe hyperparameters.",
    )
    replace_in_runs(
        p[32],
        "Validation and test descriptor targets",
        "Descriptor targets in the validation partition and locked internal test partition",
    )
    replace_in_runs(
        p[4],
        "pretraining training partition",
        "training partition",
    )
    replace_in_runs(
        p[30],
        "pretraining training partition",
        "training partition",
    )
    replace_in_runs(
        p[93],
        "pretraining training partition",
        "training partition",
    )
    replace_in_runs(
        p[98],
        "locked internal-test analysis",
        "analysis of the locked internal test partition",
    )
    replace_in_runs(
        p[107],
        "10,000-molecule pretraining-training embedding sample",
        "10,000-molecule training-partition embedding sample",
    )
    replace_in_runs(
        p[88],
        "upper bound of 50,000 validation molecules",
        "upper bound of 50,000 validation-partition molecules",
    )
    replace_in_runs(
        p[101],
        "100,000 training molecules",
        "100,000 training-partition molecules",
    )
    replace_in_runs(
        p[101],
        "10,000 training molecules",
        "10,000 training-partition molecules",
    )
    replace_in_runs(
        p[101],
        "50,000 validation embeddings",
        "50,000 validation-partition embeddings",
    )
    replace_in_runs(
        p[101],
        "5,000 validation molecules",
        "5,000 validation-partition molecules",
    )
    replace_in_runs(
        p[101],
        "The validation and training exports",
        "The validation-partition and training-partition exports",
    )
    replace_in_runs(
        p[101],
        "complete exported validation population",
        "complete exported validation-partition population",
    )
    replace_in_runs(
        p[107],
        "Evaluation used 50,000 validation embeddings",
        "Evaluation used 50,000 validation-partition embeddings",
    )
    replace_in_runs(
        p[109],
        "The primary pretraining train/validation partition was already scaffold-based",
        "The training and validation partitions were already scaffold-based",
    )
    replace_in_runs(
        p[109],
        "both pretraining training and validation data",
        "both the training and validation partitions",
    )
    replace_in_runs(
        p[109],
        "validation entries",
        "validation-partition entries",
    )
    replace_in_runs(
        p[109],
        "scaffold-disjoint validation subset/check",
        "scaffold-disjoint validation-partition subset/check",
    )
    replace_in_runs(
        p[109],
        "acyclic validation subset",
        "acyclic validation-partition subset",
    )
    replace_in_runs(
        p[112],
        "5,000-molecule validation subset",
        "5,000-molecule validation-partition subset",
    )
    replace_in_runs(
        p[116],
        "five validation molecules",
        "five validation-partition molecules",
    )
    replace_in_runs(
        p[121],
        "outer training set",
        "outer training fold",
    )
    replace_in_runs(
        p[121],
        "outer test set",
        "outer test fold",
    )
    replace_in_runs(
        p[121],
        "both partitions",
        "both outer folds",
    )
    replace_in_runs(
        p[121],
        "outer training partition",
        "outer training fold",
    )
    replace_in_runs(
        p[123],
        "complete outer training partition",
        "complete outer training fold",
    )
    replace_in_runs(
        p[123],
        "outer-training-only scaling",
        "scaling fitted only on the outer training fold",
    )
    replace_in_runs(
        p[123],
        "untouched outer test partition",
        "untouched outer test fold",
    )
    replace_in_runs(
        p[123],
        "outer-training target standard deviation",
        "outer-training-fold target standard deviation",
    )
    replace_in_runs(
        p[118],
        "five development/promotion tasks",
        "the five-dataset development/promotion panel",
    )
    replace_in_runs(
        p[118],
        "These five datasets were part of model development and fail-closed promotion",
        "These five datasets comprised the development/promotion panel and were part of fail-closed promotion",
    )
    replace_in_runs(
        p[125],
        "Outer-test metrics were",
        "Metrics on each outer test fold were",
    )
    replace_in_runs(
        p[126],
        "MoleculeNet development panel",
        "MoleculeNet development/promotion panel",
    )
    replace_in_runs(
        p[129],
        "train/test counts",
        "outer training-fold and outer test-fold counts",
    )
    replace_in_runs(
        p[133],
        "All five development datasets",
        "All five datasets in the development/promotion panel",
    )
    set_text(
        p[130],
        "Exact molecule overlap between the immutable pretraining corpus and BACE, BBBP, ESOL, "
        "FreeSolv, Lipophilicity and HIV was audited using the same downstream canonicalization "
        "and deduplication policy. Downstream canonical-isomeric-SMILES SHA-256 identities were "
        "joined to the 256 bucketed pretraining Parquet tables, and matches were counted for the "
        "complete corpus and separately for the training partition, validation partition and "
        "locked internal test partition. This audit measures molecular exposure; the experimental "
        "downstream endpoint labels were not pretraining targets.",
    )
    set_text(
        p[134],
        "1.7. Candidate checkpoints, promotion freeze, post-selection confirmation and replication",
    )
    set_text(
        p[135],
        "Retained training milestones were generated at steps 2,500, 5,000, 7,500, 10,000, "
        "12,500 and 15,000. Semantic representation selection was intentionally separate from "
        "the online reconstruction-oriented checkpoint score. Repository history shows that the "
        "criteria-bearing validator was committed and remained unchanged before the expanded "
        "complete sweep. However, the first commit already contains the criteria, the original "
        "10,000-step selection and a preliminary retained-step screen, so it cannot establish "
        "prospective specification before those analyses. Accordingly, the 17 frozen promotion "
        "criteria were applied uniformly in the complete retrospective audit of retained seed-42 "
        "checkpoints from 5,000 through 15,000 steps; the retained 2,500-step checkpoint was not "
        "included. Every candidate had its own calibrator fitted exclusively on 100,000 training-"
        "partition molecules and underwent the same validation-partition representation, "
        "similarity and clustering assessments plus the same ten-split panel on all five external "
        "development/promotion datasets. Promotion required mutually consistent checkpoint, "
        "calibrator, graph-manifest, configuration, representation-export and downstream-panel "
        "identities; all 17 criteria; and every protocol and artifact-integrity check.",
    )
    set_text(
        p[136],
        "The seed-42 checkpoint at optimization step 10,000 was the only evaluated primary-seed "
        "checkpoint to pass all 17 criteria and the complete fail-closed gate. At promotion, that "
        "checkpoint, its independently stored 100,000-molecule training-partition coordinate calibrator, "
        "the graph_z-plus-mean-node composition and the fixed mean-node weight of three were "
        "frozen. The checkpoint and calibrator were copied byte-for-byte to the canonical release "
        "names representation-best.pt and representation-calibrator.pt.",
    )
    set_text(
        p[137],
        "Only after the seed-42/10,000-step checkpoint, calibrator and 384-D embedding definition "
        "were frozen was the 1,088,766-molecule locked internal test partition opened. Its health "
        "evaluation was capped at 250,000 molecules, and a separate 50,000-molecule sample was "
        "used for geometry analyses. HIV was evaluated separately as an external post-selection "
        "confirmatory endpoint and was never part of the pretraining corpus split. The seed-43 "
        "run used independent stochastic model training and a separately fitted train-only "
        "calibrator to assess replication; it was not a seed-42 model-selection candidate. None "
        "of the locked internal test, HIV or seed-43 results could alter the promoted checkpoint, "
        "calibrator, mean-node weighting or promotion decision.",
    )

    insert_evidence_roles_table(document, p[138])
    set_text(p[138], "1.9. Standalone inference and artifact-integrity checks")
    set_text(p[141], "1.10. Software environment and computational reproducibility")
    set_text(p[144], "1.11. Code and data availability")

    # Results introduction answers the evidence-role questions before details.
    set_text(
        p[147],
        "Results are organized by evidence source. Encoder weights were learned only from the "
        "221,148,895-graph training partition; the selected seed-42/10,000-step checkpoint had "
        "actually consumed 57,504,265 unique training graphs. The 943,038-graph validation "
        "partition and the five external development/promotion datasets could influence "
        "checkpoint promotion but never neural-network weights. After the checkpoint, calibrator "
        "and 384-D representation were frozen, the 1,088,766-graph locked internal test partition "
        "was opened for protected post-selection evaluation. HIV supplied a separate external "
        "post-selection endpoint, and seed 43 supplied independent training replication rather "
        "than model selection. For readability, the subsections below are grouped by evidence "
        "source while retaining this chronology.",
    )

    # Split the old mixed architecture/seed-43 paragraph without dropping content.
    architecture_summary = new_paragraph_before(
        p[195],
        "The v5 architecture replaced the stochastic VGAE latent with deterministic "
        "128-dimensional atom embeddings and an explicit 256-dimensional graph vector generated "
        "by four residual GINE blocks. Masked atom and bond reconstruction, dropped-bond recovery, "
        "model-hard non-edge discrimination, 13 auxiliary descriptor targets and a weak mean-node "
        "NT-Xent term jointly supplied self-supervision.",
        "Normal",
    )
    set_text(
        p[153],
        "Unlike the earlier models, the independent seed-43 encoder showed no late-collapse "
        "trajectory. All 256 graph coordinates remained active throughout 15,000 steps; clean-"
        "graph effective rank increased from 26.51 at step 2,500 to 37.45 at step 10,000 and "
        "39.12 at step 15,000, while median coordinate standard deviation remained within "
        "0.72-0.77. These results were replication evidence and could not influence selection of "
        "the seed-42 checkpoint.",
    )

    # Results headings.
    for paragraph, title in zip(
        (p[149], p[152], p[155], p[162], p[176], p[186], p[191], p[170]),
        RESULT_HEADINGS,
    ):
        set_text(paragraph, title)

    # 2.1.1: split counts and exact exposure, with the training-progress figure.
    replace_in_runs(
        p[150],
        "943,038 in validation and",
        "943,038 in the validation partition and",
    )
    replace_in_runs(
        p[150],
        "1,088,766 in the locked internal test.",
        "1,088,766 in the locked internal test partition.",
    )
    replace_in_runs(
        p[150],
        "Training was deliberately step-based",
        "Training and validation progress are summarized in Fig. 1. Training was deliberately step-based",
    )

    # 2.1.2: validation sample sizes and provenance made explicit.
    replace_in_runs(
        p[154],
        "evaluation on 50,000 validation graphs",
        "evaluation on 50,000 graphs sampled from the 943,038-graph validation partition",
    )
    replace_in_runs(
        p[164],
        "authoritative 50,000-molecule validation panel",
        "authoritative 50,000-molecule validation-partition panel",
    )
    set_text(
        p[166],
        "To test whether the representation encoded information beyond the 13 descriptors used "
        "as auxiliary training targets, a frozen ridge-regression probe was fitted on 10,000 "
        "training-partition embeddings and evaluated on 50,000 validation-partition embeddings "
        "for 13 additional topological quantities absent from the pretraining objective. Across "
        "these targets, the representation achieved mean R² 0.9679, median R² 0.9675 and mean "
        "standardized absolute error 0.1120. A descriptor-only control under the same probe "
        "framework achieved mean R² 0.7701, median R² 0.8584 and mean standardized absolute error "
        "0.3229; the embedding therefore increased mean R² by 0.1978. On the explicitly filtered "
        "49,047-molecule non-empty-scaffold validation-partition subset, embedding mean R² remained 0.9681. "
        "The weakest individual quantity was the more demanding Kier kappa3 shape index (R² "
        "approximately 0.90), whereas several other quantities, including fraction sp3 and "
        "multiple connectivity/shape descriptors, approached or exceeded R² of 0.98.",
    )
    set_text(
        p[168],
        "The scaffold-disjoint topology value should not be interpreted as evidence from an "
        "additional, stricter scaffold split. The training and validation partitions "
        "were already scaffold-separated, so validation-partition molecules with non-empty "
        "Bemis-Murcko scaffolds were, by construction, separated from the full training partition "
        "at scaffold level. The probe-level filter retained validation-partition molecules only when the "
        "stored scaffold was non-empty and absent from the finite 10,000-molecule probe-training "
        "export. For scaffolded molecules this condition therefore largely reiterates the global "
        "split, while acyclic molecules with an empty Murcko scaffold are excluded. The near-"
        "identical mean R² for the full validation-partition panel (0.9679) and this subset (0.9681) is best "
        "interpreted as a robustness check showing that the topology result is not driven by the "
        "acyclic validation-partition subset, rather than as a second escalation of scaffold novelty.",
    )
    replace_in_runs(
        p[171],
        "The validation probes instead showed",
        "Within a 5,000-molecule similarity subset drawn from the authoritative 50,000-molecule validation-partition export, the probes showed",
    )
    set_text(
        p[173],
        "Unsupervised analysis using standard Euclidean K-means on L2-normalized inputs provided "
        "a complementary global view; because the fitted centroids were not renormalized, this "
        "was not spherical K-means. From the authoritative 50,000-molecule validation-partition "
        "export, 11,034 molecules belonging to recurring Bemis-Murcko scaffolds entered the "
        "clustering analysis. gMolAI achieved adjusted Rand index (ARI) 0.362 and normalized "
        "mutual information (NMI) 0.788, compared with ARI 0.317 and NMI 0.777 for matched Morgan "
        "fingerprints. A second independently sampled 50,000-molecule validation-partition panel "
        "contributed 11,011 recurring-scaffold molecules and again favored gMolAI numerically "
        "(ARI 0.352 versus 0.322; NMI 0.785 versus 0.783). These results indicate that scaffold-"
        "related organization can emerge from the learned continuous space without requiring it "
        "to become an ECFP replica.",
    )

    # 2.1.3: frozen retrospective chronology and shifted cross-references.
    set_text(
        p[156],
        "Complete checkpoint evaluation revealed a key decoupling between continued pretraining "
        "improvement and downstream eligibility. The 17 frozen promotion criteria were applied "
        "uniformly in the complete retrospective 5,000-15,000-step seed-42 audit. Promotion "
        "combined representation-validation evidence with the same five external development/"
        "promotion datasets, using a separate training-only calibrator at every checkpoint. Step "
        "5,000 passed 15 of 17 criteria, failing effective rank (24.865 versus the minimum of 25) "
        "and FreeSolv RMSE (1.39045 versus the maximum of 1.30). Steps 7,500, 12,500 and 15,000 "
        "each passed 16 of 17 and failed only FreeSolv, with RMSE 1.37227, 1.32289 and 1.30346, "
        "respectively. Step 10,000 passed all 17 criteria, including FreeSolv RMSE 1.29716, and "
        "was the only checkpoint to pass the complete fail-closed gate (Table 8; Fig. 2). "
        "Protocol, identity, source-integrity and artifact-completeness checks passed at every "
        "evaluated step. No threshold was modified during this expanded audit, and the retained "
        "2,500-step checkpoint was not included.",
    )

    # 2.1.4: keep the representation definition concise and normalize terminology.
    replace_in_runs(
        p[163],
        "100,000 pretraining-training molecules",
        "100,000 training-partition molecules",
    )

    # 2.1.5: development evidence, descriptor control and exposure are explicitly selection-conditioned.
    replace_in_runs(
        p[177],
        "five MoleculeNet development/promotion datasets,",
        "the five MoleculeNet development/promotion datasets,",
    )
    replace_in_runs(
        p[177],
        "HIV and the locked internal test were reserved for post-selection evaluation. ",
        "These outer test folds were benchmark folds and were not the locked internal test partition. ",
    )
    set_text(
        p[178],
        "Exact canonical-identity auditing distinguished molecules present somewhere in the "
        "pretraining corpus, molecules assigned to its training partition and molecules actually "
        "consumed before each checkpoint. Complete-corpus overlaps were 414/1,513 BACE, "
        "1,090/1,860 BBBP, 969/1,116 ESOL, 526/639 FreeSolv and 2,513/4,198 Lipophilicity; "
        "training-partition memberships were 413, 1,080, 964, 524 and 2,493, respectively. At the "
        "promoted step 10,000, only 109 BACE, 278 BBBP, 226 ESOL, 120 FreeSolv and 682 "
        "Lipophilicity training-overlap molecules had actually been consumed, corresponding to "
        "26.39%, 25.74%, 23.44%, 22.90% and 27.36% of their training-partition overlaps. By step "
        "15,000 the corresponding consumed counts were 161, 413, 371, 186 and 989. Every five-"
        "checkpoint trajectory was monotonic; Table 11 reports both percent of the full downstream "
        "dataset and percent of training-partition overlaps. Thus, training-partition membership "
        "is not equivalent to actual pre-checkpoint consumption. These molecules carried graph-"
        "derived and calculated-descriptor pretraining targets, not experimental endpoint labels; "
        "this is not direct endpoint-label leakage, but actual consumption limits molecule-level "
        "novelty claims for the affected subset.",
    )
    replace_in_runs(p[179], "(Table 9)", "(Table 10)")

    # 2.1.6: seed 43 is replication only.
    replace_in_runs(
        p[187],
        "The complete model-training experiment was independently repeated with seed 43.",
        "The complete model-training experiment was independently repeated with seed 43 as replication evidence, not as a candidate for seed-42 model selection.",
    )

    # 2.1.7: protected internal evidence, two samples, and the clustering reversal.
    set_text(
        p[193],
        "The 1,088,766-molecule locked internal test partition remained unopened until the "
        "seed-42/10,000-step checkpoint, calibrator and 384-D representation were frozen. The "
        "latent/reconstruction health analysis used the permitted maximum of 250,000 molecules: "
        "all 256 global coordinates remained active, clean effective rank was 37.68, median "
        "coordinate standard deviation was 0.732, and model-hard edge discrimination achieved "
        "AUROC 0.9661 and average precision 0.8286 without non-finite values. A separate 50,000-"
        "molecule geometry panel drawn from the same locked internal test partition yielded "
        "calibrated-embedding effective rank 30.69, Morgan recall@10 0.2079, cosine-Tanimoto "
        "Spearman correlation 0.4573, neighbour Tanimoto enrichment 2.00-fold and scaffold-"
        "neighbour enrichment 26.36-fold. These values closely reproduced the validation-"
        "partition retrieval phenotype.",
    )
    set_text(
        p[194],
        "The scientifically important reversal occurred in the 16,382-molecule recurring-"
        "scaffold subset of the geometry panel from the locked internal test partition: Morgan outperformed gMolAI "
        "for scaffold clustering, with ARI/NMI 0.372/0.780 versus 0.350/0.722 for gMolAI. The "
        "gMolAI clustering advantage on the validation partition therefore did not generalize as "
        "a universal superiority under a different scaffold-hash population. Morgan fingerprints "
        "are explicitly constructed from local circular substructures and can be particularly "
        "well matched to scaffold-oriented grouping. The value of gMolAI is instead supported by "
        "the broader conjunction of stable topology accessibility, chemically enriched retrieval, "
        "external task transfer, cross-seed reproducibility and a distinct continuous geometry. "
        "No result from the locked internal test partition was allowed to change the checkpoint, calibrator, mean-node "
        "weighting or promotion decision.",
    )

    # 2.1.8: HIV is external and post-selection, not part of the internal corpus split.
    set_text(
        p[192],
        "HIV was a separate external MoleculeNet dataset, never part of the training partition, "
        "validation partition or locked internal test partition. It was evaluated only after seed-42 "
        "selection and could not influence promotion. After chemical filtering, HIV contained "
        "37,225 molecules in 18,651 scaffolds, with 3.092% positives. Exact identity auditing "
        "found 27,377 molecules present somewhere in the pretraining corpus and 27,145 assigned "
        "to its training partition; 7,038 had actually been consumed by step 10,000 (18.91% of "
        "the full HIV dataset and 25.93% of training-partition overlaps; Table 11). Across ten "
        "nested scaffold splits, gMolAI achieved ROC-AUC 0.7578 ± 0.0164, average precision "
        "0.1813 ± 0.0202 and balanced accuracy 0.6856 ± 0.0241; Morgan achieved ROC-AUC "
        "0.7440 ± 0.0258, average precision 0.1945 ± 0.0417 and balanced accuracy "
        "0.6830 ± 0.0190. Values after ± are descriptive population s.d. (ddof=0). gMolAI was "
        "numerically higher in ROC-AUC and balanced accuracy, whereas Morgan retained higher "
        "average precision on this strongly imbalanced confirmatory endpoint. Repeated outer "
        "scaffold splits can overlap, so these dispersions are not standard errors from independent "
        "replicates; the result does not support universal superiority over Morgan.",
    )

    # Renumber existing result tables after inserting the new Methods Table 6.
    shifted_captions = {
        6: "Table 8 | Complete fail-closed promotion evaluation across retained primary-seed checkpoints.",
        7: "Table 9 | Pretraining diagnostics continued to improve after the selected checkpoint whereas FreeSolv transfer did not.",
        8: "Table 7 | Validation-partition representation probes for the selected seed-42 checkpoint.",
        9: "Table 10 | Selection-conditioned external development/promotion benchmarks for the selected 384-D representation, Morgan fingerprints and the 13-descriptor-only control.",
        10: "Table 11 | Exact checkpoint-resolved downstream-molecule exposure in the seed-42 training stream, with HIV shown as confirmatory context.",
        11: "Table 12 | Independent seed-43 training replication at step 10,000.",
    }
    for old_number, text in shifted_captions.items():
        caption = find_exact_paragraph(document, EXPECTED_REV4_TABLE_CAPTIONS[old_number - 1])
        set_text(caption, text, bold=True)
        caption.paragraph_format.keep_with_next = True
    replace_in_runs(
        p[175],
        "All gates were evaluated",
        "The main metrics used a 50,000-molecule validation-partition export; similarity used 5,000 molecules, and recurring-scaffold clustering used 11,034 molecules. A separate 50,000-molecule validation-partition panel contributed 11,011 clustering molecules. All gates were evaluated",
    )
    replace_in_runs(
        p[175],
        "global pretraining train/validation split was already scaffold-based",
        "global training and validation partitions were already scaffold-based",
    )
    replace_in_runs(
        p[175],
        "non-empty-scaffold validation subset/check",
        "non-empty-scaffold validation-partition subset/check",
    )
    replace_in_runs(
        p[188],
        "Downstream development-panel results",
        "Downstream development/promotion-panel results",
    )
    replace_in_runs(
        p[190],
        "pretraining training sample",
        "training-partition sample",
    )
    p[185].style = "Table Note"
    set_text(
        p[185],
        "Corpus and Train cells give n and % of the post-filter dataset. Checkpoint cells give "
        "the exact number consumed before the checkpoint and, on the second line, % of the full "
        "downstream dataset/% of training-partition overlaps. *HIV is included only as external "
        "confirmatory context and was not used for checkpoint promotion.",
    )

    # Discussion cross-reference and evidence-role interpretation.
    replace_in_runs(p[205], "reported in Table 10", "reported in Table 11")
    set_text(
        p[210],
        "Taken together, the results support a deliberately bounded conclusion. The selected "
        "seed-42 gMolAI-v2.0 encoder learns a stable continuous molecular representation that "
        "avoids the collapse observed in earlier implementations, retains scaffold-general "
        "structural information, defines a chemically meaningful but non-Morgan-equivalent "
        "geometry and uniquely passed the 17 frozen promotion criteria in the complete "
        "retrospective checkpoint audit. The five external development/promotion datasets are "
        "selection-conditioned; seed 43 independently replicated the representation and transfer "
        "phenotype but did not select the released model. HIV provides a separate external post-"
        "selection endpoint, while the locked internal test partition provides protected internal "
        "confirmation and showed no numerical instability or latent collapse. Morgan remained "
        "stronger for some endpoints and, importantly, for scaffold clustering on the locked "
        "internal test partition. The evidence therefore supports complementarity rather than "
        "replacement, within the limits imposed by auxiliary descriptor supervision and measured "
        "pretraining exposure of some downstream molecules.",
    )

    # Remove now-redundant nested Results headings; all substantive paragraphs are retained.
    delete_paragraph(p[165])
    delete_paragraph(p[167])

    # Reassemble Results in evidence-source order, keeping each table beside its interpretation.
    ordered_elements = [
        p[149]._p,
        p[150]._p,
        architecture_summary._p,
        p[151]._p,
        p[152]._p,
        p[154]._p,
        p[164]._p,
        p[166]._p,
        p[168]._p,
        p[169]._p,
        p[171]._p,
        p[172]._p,
        p[173]._p,
        p[174]._p,
        old_result_tables[8],
        p[175]._p,
        p[155]._p,
        p[156]._p,
        p[157]._p,
        old_result_tables[6],
        p[158]._p,
        p[159]._p,
        p[160]._p,
        old_result_tables[7],
        p[161]._p,
        p[162]._p,
        p[163]._p,
        p[176]._p,
        p[177]._p,
        p[179]._p,
        p[180]._p,
        old_result_tables[9],
        p[181]._p,
        p[182]._p,
        p[183]._p,
        p[178]._p,
        p[184]._p,
        old_result_tables[10],
        p[185]._p,
        p[186]._p,
        p[153]._p,
        p[187]._p,
        p[188]._p,
        p[189]._p,
        old_result_tables[11],
        p[190]._p,
        p[191]._p,
        p[193]._p,
        p[194]._p,
        p[170]._p,
        p[192]._p,
    ]
    for element in ordered_elements:
        p[195]._p.addprevious(element)

    # Fail-closed structural and scientific-content checks before serialization.
    headings = tuple(
        paragraph.text
        for paragraph in document.paragraphs
        if re.match(r"^2\.1\.[1-8]\. ", paragraph.text)
    )
    if headings != RESULT_HEADINGS:
        raise RuntimeError(f"Unexpected Results headings/order: {headings!r}")
    captions = [
        paragraph.text
        for paragraph in document.paragraphs
        if re.match(r"^Table \d+ \|", paragraph.text)
    ]
    caption_numbers = [int(re.match(r"^Table (\d+) \|", text).group(1)) for text in captions]
    if caption_numbers != list(range(1, 13)):
        raise RuntimeError(f"Non-sequential table captions: {caption_numbers}")
    if len(document.tables) != 12:
        raise RuntimeError(f"Expected 12 rev5 tables, found {len(document.tables)}")
    new_matrices = [table_matrix(table) for table in document.tables]
    expected_result_matrices = [original_tables[index] for index in (7, 5, 6, 8, 9, 10)]
    if (new_matrices[:5] != original_tables[:5]
            or new_matrices[6:] != expected_result_matrices):
        raise RuntimeError("An existing scientific result table changed during restructuring")
    if omml_hashes(document) != original_omml:
        raise RuntimeError("An existing equation changed during restructuring")

    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required_literals = (
        "223,180,699",
        "221,148,895",
        "943,038",
        "1,088,766",
        "28,743,683",
        "43,109,793",
        "57,504,265",
        "71,870,280",
        "86,236,032",
        "10,000 passed all 17",
        "2,500-step checkpoint was not included",
        "not spherical K-means",
        "approximately 80%/20% of scaffold groups",
        "Morgan outperformed gMolAI",
        "No result from the locked internal test partition was allowed to change",
    )
    missing = [literal for literal in required_literals if literal not in full_text]
    if missing:
        raise RuntimeError(f"Required scientific/interpretive statements missing: {missing}")
    forbidden = (
        "prospectively pre-specified",
        "prospectively prespecified",
        "structurally unseen",
        "validation set",
        "validation-set",
        "internal test set",
    )
    present = [term for term in forbidden if term.lower() in full_text.lower()]
    if present:
        raise RuntimeError(f"Ambiguous or unsupported terminology remains: {present}")
    figure_mentions = re.findall(r"Fig\.\s*(\d+)", full_text)
    if figure_mentions != ["1", "2"]:
        raise RuntimeError(f"Unexpected figure-reference sequence: {figure_mentions}")

    document.core_properties.subject = "gMolAI-v2.0 manuscript revision 5"
    document.core_properties.revision = 5
    destination.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary = tempfile.mkstemp(
        prefix=f".{destination.name}.", suffix=".tmp", dir=destination.parent
    )
    os.close(fd)
    try:
        document.save(temporary)
        os.chmod(temporary, source.stat().st_mode & 0o777)
        os.replace(temporary, destination)
    except BaseException:
        try:
            os.unlink(temporary)
        except FileNotFoundError:
            pass
        raise

    # Reopen the finished package and repeat the table/equation invariants.
    reopened = Document(destination)
    if len(reopened.tables) != 12:
        raise RuntimeError("Serialized rev5 did not reopen with 12 tables")
    reopened_matrices = [table_matrix(table) for table in reopened.tables]
    if (reopened_matrices[:5] != original_tables[:5]
            or reopened_matrices[6:] != expected_result_matrices):
        raise RuntimeError("Serialized rev5 changed an existing table")
    if omml_hashes(reopened) != original_omml:
        raise RuntimeError("Serialized rev5 changed an existing equation")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    return parser.parse_args()


if __name__ == "__main__":
    build(parse_args())
