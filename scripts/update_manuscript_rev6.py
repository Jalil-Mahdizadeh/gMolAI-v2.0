#!/usr/bin/env python3
"""Create gMolAI manuscript revision 6 from the authoritative revision 5.

Revision 6 adds only the completed frozen-encoder locked-test comparison: a
common-panel radar figure, a compact exact-parameter table, and a bounded
explanation of the single pathological RDKit Kappa3 outlier.  It does not load
or execute a molecular encoder and it does not alter any existing result.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import os
import re
import tempfile
from collections import Counter
from pathlib import Path

from lxml import etree

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.text.paragraph import Paragraph
except ImportError as error:  # pragma: no cover - build dependency
    raise SystemExit("python-docx==1.2.0 is required to build manuscript rev6") from error


AUTHORITATIVE_REV5_SHA256 = (
    "9a17e5384a61212c8c462b3f6a267ece8a6bc53dd4370c13afa4c46462a6dbb2"
)
BENCHMARK_SUMMARY_SHA256 = (
    "e5946cf9b577e9906c810ad810eb3af59cb76a0ac0dea274bb77acb5d3b2812c"
)
PARAMETER_COUNTS_SHA256 = (
    "67cb943475dea1ef1454142ff0ce8027d8243ab5ba79f6de6eab7f8612b46e82"
)
FIGURE3_PNG_SHA256 = (
    "1f16faef61ccae1065a6bc7bd2ba35e77b14d88fcdf1b57b1dcc7e4ad037d5f5"
)
MODEL_ORDER = (
    "gmolai",
    "morgan",
    "molai",
    "molformer",
    "smi_ted",
    "molclr_gin",
    "kermt_v2",
)
MODEL_LABELS = {
    "gmolai": "gMolAI",
    "morgan": "Morgan radius-2",
    "molai": "MolAI epoch 6",
    "molformer": "MoLFormer",
    "smi_ted": "SMI-TED-Light",
    "molclr_gin": "MolCLR-GIN",
    "kermt_v2": "KERMT v2",
}
EXPECTED_TOTAL_PARAMETERS = {
    "gmolai": 2_220_753,
    "morgan": 0,
    "molai": 24_281_600,
    "molformer": 44_375_040,
    "smi_ted": 289_913_857,
    "molclr_gin": 2_404_196,
    "kermt_v2": 49_209_854,
}
EXPECTED_ENCODING_PARAMETERS = {
    "gmolai": 1_445_252,
    "morgan": 0,
    "molai": 24_281_600,
    "molformer": 44_375_040,
    "smi_ted": 166_565_376,
    "molclr_gin": 2_404_196,
    "kermt_v2": 49_209_854,
}
EXPECTED_DIMENSIONS = {
    "gmolai": 384,
    "morgan": 2_048,
    "molai": 512,
    "molformer": 768,
    "smi_ted": 768,
    "molclr_gin": 512,
    "kermt_v2": 512,
}
EXPECTED_METRICS = {
    "gmolai": {
        "topology_mean_r2": 0.9705115056488679,
        "scaffold_disjoint_mean_r2": 0.9728505497473741,
        "clustering_ari": 0.35776229752630034,
        "clustering_nmi": 0.724601855860231,
    },
    "morgan": {
        "clustering_ari": 0.3955703042792596,
        "clustering_nmi": 0.7927995386560196,
    },
    "kermt_v2": {
        "clustering_ari": 0.3619956070084435,
        "clustering_nmi": 0.754795257961806,
    },
}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def require_hash(path: Path, expected: str, label: str) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)
    observed = sha256_file(path)
    if observed != expected:
        raise RuntimeError(f"{label} hash changed: {observed}; expected {expected}")


def find_exact_paragraph(document, text: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact paragraph {text!r}, found {len(matches)}")
    return matches[0]


def find_paragraph(document, prefix: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}, found {len(matches)}")
    return matches[0]


def paragraph_after(anchor: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def paragraph_before(anchor: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def table_matrix(table) -> tuple[tuple[str, ...], ...]:
    return tuple(tuple(cell.text for cell in row.cells) for row in table.rows)


def omml_hashes(document) -> Counter[str]:
    hashes: Counter[str] = Counter()
    for tag in (qn("m:oMath"), qn("m:oMathPara")):
        for node in document.element.body.iter(tag):
            hashes[hashlib.sha256(etree.tostring(node)).hexdigest()] += 1
    return hashes


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_cant_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def set_cell(cell, text: str, *, bold: bool = False, size: float = 8.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    properties = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", 30), ("left", 45), ("bottom", 30), ("right", 45)):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    properties.append(margins)


def load_benchmark(path: Path) -> dict[str, dict[str, str]]:
    require_hash(path, BENCHMARK_SUMMARY_SHA256, "benchmark summary")
    with path.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))
    by_model = {row["model"]: row for row in rows}
    if tuple(row["model"] for row in rows) != MODEL_ORDER or len(by_model) != 7:
        raise RuntimeError("Benchmark summary model order or uniqueness changed")
    for model, expected_metrics in EXPECTED_METRICS.items():
        for metric, expected in expected_metrics.items():
            observed = float(by_model[model][metric])
            if observed != expected:
                raise RuntimeError(f"{model} {metric} changed: {observed} != {expected}")
    if {int(row["train_rows"]) for row in rows} != {9_958}:
        raise RuntimeError("Common topology-probe training count changed")
    if {int(row["test_rows"]) for row in rows} != {49_844}:
        raise RuntimeError("Common locked-test count changed")
    if {int(row["clustering_rows"]) for row in rows} != {16_360}:
        raise RuntimeError("Common recurring-scaffold clustering count changed")
    return by_model


def load_parameter_counts(path: Path) -> dict[str, dict[str, str]]:
    require_hash(path, PARAMETER_COUNTS_SHA256, "parameter-count CSV")
    with path.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))
    by_model = {row["benchmark_key"]: row for row in rows}
    if tuple(row["benchmark_key"] for row in rows) != MODEL_ORDER or len(by_model) != 7:
        raise RuntimeError("Parameter-count model order or uniqueness changed")
    for model in MODEL_ORDER:
        if int(by_model[model]["total_loaded_parameters"]) != EXPECTED_TOTAL_PARAMETERS[model]:
            raise RuntimeError(f"Total parameter count changed for {model}")
        if int(by_model[model]["embedding_path_parameters"]) != EXPECTED_ENCODING_PARAMETERS[model]:
            raise RuntimeError(f"Encoding-path parameter count changed for {model}")
        if int(by_model[model]["embedding_dimension"]) != EXPECTED_DIMENSIONS[model]:
            raise RuntimeError(f"Embedding dimension changed for {model}")
    return by_model


def old_text_is_subsequence(old: list[str], new: list[str]) -> bool:
    position = 0
    for text in new:
        if position < len(old) and text == old[position]:
            position += 1
    return position == len(old)


def insert_parameter_table(document, anchor: Paragraph, counts: dict[str, dict[str, str]]):
    caption = paragraph_before(
        anchor,
        "Table 13 | Exact parameter counts for frozen representations in the common-panel benchmark.",
        "Table Caption",
    )
    caption.runs[0].bold = True
    caption.paragraph_format.keep_with_next = True

    reference_style = document.tables[-1].style
    table = document.add_table(rows=1, cols=3)
    table.style = reference_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    anchor._p.addprevious(table._tbl)
    headers = ("Representation", "Output dimensions", "Total loaded parameters")
    widths = (2.65, 1.35, 2.25)
    for index, (header, width) in enumerate(zip(headers, widths)):
        set_cell(table.rows[0].cells[index], header, bold=True)
        table.rows[0].cells[index].width = Inches(width)
    set_repeat_header(table.rows[0])
    set_cant_split(table.rows[0])
    for model in MODEL_ORDER:
        row = table.add_row()
        set_cant_split(row)
        values = (
            MODEL_LABELS[model],
            f"{int(counts[model]['embedding_dimension']):,}",
            f"{int(counts[model]['total_loaded_parameters']):,}",
        )
        for index, (value, width) in enumerate(zip(values, widths)):
            set_cell(row.cells[index], value)
            row.cells[index].width = Inches(width)
    return table


def build(args: argparse.Namespace) -> None:
    source = Path(args.input).resolve()
    destination = Path(args.output).resolve()
    benchmark_path = Path(args.benchmark_summary).resolve()
    parameter_path = Path(args.parameter_counts).resolve()
    figure_path = Path(args.figure).resolve()
    if source == destination:
        raise ValueError("rev6 output must differ from the retained rev5 input")
    require_hash(source, AUTHORITATIVE_REV5_SHA256, "authoritative rev5")
    require_hash(figure_path, FIGURE3_PNG_SHA256, "Figure 3 PNG")
    benchmark = load_benchmark(benchmark_path)
    counts = load_parameter_counts(parameter_path)

    document = Document(source)
    if len(document.tables) != 12:
        raise RuntimeError(f"Expected 12 rev5 tables, found {len(document.tables)}")
    if len(document.inline_shapes) != 0:
        raise RuntimeError("Authoritative rev5 unexpectedly contains inline figures")
    old_tables = [table_matrix(table) for table in document.tables]
    old_omml = omml_hashes(document)
    old_paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
    old_caption_numbers = [
        int(match.group(1))
        for paragraph in document.paragraphs
        if (match := re.match(r"^Table (\d+) \|", paragraph.text))
    ]
    if old_caption_numbers != list(range(1, 13)):
        raise RuntimeError(f"Unexpected rev5 table captions: {old_caption_numbers}")

    methods_anchor = find_paragraph(
        document,
        "Only after the seed-42/10,000-step checkpoint, calibrator and 384-D embedding definition",
    )
    paragraph_after(
        methods_anchor,
        "As a retrospective post-selection comparison, the unchanged representation probes were "
        "applied to gMolAI, Morgan radius-2 fingerprints, MolAI epoch 6, MoLFormer, SMI-TED-"
        "Light, MolCLR-GIN and KERMT v2. Coverage was screened before comparison, and every "
        "reported cross-model metric used the identical ordered intersection: 9,958 training-"
        "partition molecules for fitting the topology probe and 49,844 molecules from the locked "
        "internal test panel for evaluation. No endpoint labels were available or used. Exact "
        "parameter totals were obtained by instantiating the frozen benchmark modules in their "
        "pinned Apptainer images and summing PyTorch parameter elements; Morgan is algorithmic "
        "and therefore has no learned parameters. This additive analysis could not alter model "
        "selection or any released artifact.",
        "Body Text",
    )

    results_anchor = find_exact_paragraph(
        document, "2.1.8. HIV — independent external post-selection endpoint"
    )
    paragraph_before(
        results_anchor,
        "The post-selection frozen-encoder comparison used 9,958 common probe-training molecules "
        "and 49,844 common molecules from the locked internal test panel. Fig. 3 presents six "
        "complementary metrics, each divided by the highest observed value on that spoke; this "
        "visual normalization defines neither an aggregate score nor a significance test. gMolAI "
        "had the highest topology mean R² (0.9705) and scaffold-disjoint topology mean R² "
        "(0.9729). Morgan had the strongest recurring-scaffold clustering (ARI 0.3956; NMI "
        "0.7928), while KERMT v2 was the closest learned comparator for clustering (ARI 0.3620; "
        "NMI 0.7548). Thus, the matched comparison supports complementary strengths rather than "
        "universal superiority.",
        "Normal",
    )
    picture = paragraph_before(results_anchor, "", "Normal")
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_with_next = True
    shape = picture.add_run().add_picture(str(figure_path), width=Inches(6.20))
    shape._inline.docPr.set(
        "title", "Frozen-encoder comparison on the common locked-test panel"
    )
    shape._inline.docPr.set(
        "descr",
        "Radar plot comparing seven frozen molecular representations across topology, neighbour "
        "enrichment, scaffold-neighbour purity and scaffold-clustering metrics. Each spoke is "
        "normalized to the best observed value; gMolAI leads topology and Morgan leads clustering.",
    )
    figure_caption = paragraph_before(results_anchor, "", "Caption")
    figure_caption.paragraph_format.keep_with_next = False
    label = figure_caption.add_run("Figure 3 | ")
    label.bold = True
    figure_caption.add_run(
        "Frozen-encoder comparison on the all-model common locked internal-test panel. Six raw "
        "metrics were independently divided by their maximum across the seven representations so "
        "that the best observed value on each spoke equals 1.0. This best-relative display does "
        "not define a composite score. Topology probes used 9,958 common training-partition "
        "molecules and 49,844 common locked-test molecules; clustering used 16,360 recurring-"
        "scaffold molecules. Exact raw values are retained in the source artifact."
    )

    scale_paragraph = paragraph_before(
        results_anchor,
        "Model scale differed by more than two orders of magnitude (Table 13). The table reports "
        "the literal loaded-module count requested for reproducibility; this includes auxiliary "
        "heads in gMolAI and decoder components in SMI-TED-Light even when they do not contribute "
        "to the exported vector.",
        "Normal",
    )
    scale_paragraph.paragraph_format.page_break_before = True
    inserted_table = insert_parameter_table(document, results_anchor, counts)
    table_note = paragraph_before(
        results_anchor,
        "Counts are exact parameter elements, not trainable parameters during this frozen "
        "benchmark. The gMolAI and SMI-TED-Light encoding paths contain 1,445,252 and "
        "166,565,376 unique parameters, respectively, versus loaded totals of 2,220,753 and "
        "289,913,857. Morgan fingerprints contain no learned parameters.",
        "Table Note",
    )
    inserted_table._tbl.addnext(table_note._p)

    paragraph_before(
        results_anchor,
        "The apparent increase in gMolAI topology mean R² from 0.9018 on the complete "
        "10,000/50,000 panels to 0.9705 on the common panels was traced to one pathological "
        "descriptor row rather than a general performance change. MolAI's tokenizer excluded 42 "
        "training and 156 test structures containing unsupported 'B' and/or lowercase 'i'; among "
        "them was the acyclic silicon molecule FS(F)(F)(F)(F)[SiH3], because the token 'Si' "
        "contains lowercase 'i'. RDKit assigned that molecule Kappa3 = 2350.1974, whereas the fitted probe predicted "
        "approximately 2.41. This single molecule contributed 99.831% of the complete-panel "
        "Kappa3 residual sum of squares; removing only it changed Kappa3 R² from 0.0126 to 0.8933. "
        "The scaffold-disjoint topology mean remained stable (0.97226 complete versus 0.97285 "
        "common) because empty-scaffold molecules were already excluded. Accordingly, the common "
        "panel is the primary matched-model comparison, while the complete-panel result is "
        "retained as an explicit sensitivity analysis.",
        "Normal",
    )

    if [table_matrix(table) for table in document.tables[:12]] != old_tables:
        raise RuntimeError("An existing rev5 scientific table changed")
    if omml_hashes(document) != old_omml:
        raise RuntimeError("An existing rev5 equation changed")
    if not old_text_is_subsequence(old_paragraph_texts, [p.text for p in document.paragraphs]):
        raise RuntimeError("Existing rev5 paragraph text changed or moved")
    captions = [
        int(match.group(1))
        for paragraph in document.paragraphs
        if (match := re.match(r"^Table (\d+) \|", paragraph.text))
    ]
    if captions != list(range(1, 14)):
        raise RuntimeError(f"Non-sequential rev6 table captions: {captions}")
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    figure_mentions = re.findall(r"Fig\.\s*(\d+)", full_text)
    if figure_mentions != ["1", "2", "3"]:
        raise RuntimeError(f"Unexpected in-text figure-reference sequence: {figure_mentions}")
    required = (
        "9,958 common probe-training molecules",
        "49,844 common molecules",
        "0.9705",
        "0.9729",
        "0.3956",
        "0.7928",
        "FS(F)(F)(F)(F)[SiH3]",
        "Kappa3 = 2350.1974",
        "99.831%",
        "0.0126 to 0.8933",
        "0.97226 complete versus 0.97285 common",
        "Table 13",
    )
    missing = [value for value in required if value not in full_text]
    if missing:
        raise RuntimeError(f"Required rev6 content is missing: {missing}")
    if len(document.tables) != 13 or len(document.inline_shapes) != 1:
        raise RuntimeError("rev6 must contain 13 tables and one embedded radar figure")

    document.core_properties.subject = "gMolAI-v2.0 manuscript revision 6"
    document.core_properties.revision = 6
    destination.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary = tempfile.mkstemp(
        prefix=f".{destination.name}.", suffix=".tmp", dir=destination.parent
    )
    os.close(fd)
    try:
        document.save(temporary)
        os.chmod(temporary, source.stat().st_mode & 0o777)
        os.replace(temporary, destination)
    except BaseException:
        try:
            os.unlink(temporary)
        except FileNotFoundError:
            pass
        raise

    reopened = Document(destination)
    if len(reopened.tables) != 13 or len(reopened.inline_shapes) != 1:
        raise RuntimeError("Serialized rev6 did not reopen with 13 tables and one figure")
    if [table_matrix(table) for table in reopened.tables[:12]] != old_tables:
        raise RuntimeError("Serialized rev6 changed an existing table")
    if table_matrix(reopened.tables[12]) != table_matrix(inserted_table):
        raise RuntimeError("Serialized rev6 parameter table changed")
    if omml_hashes(reopened) != old_omml:
        raise RuntimeError("Serialized rev6 changed an existing equation")
    if not old_text_is_subsequence(old_paragraph_texts, [p.text for p in reopened.paragraphs]):
        raise RuntimeError("Serialized rev6 changed existing paragraph text or order")
    if sha256_file(source) != AUTHORITATIVE_REV5_SHA256:
        raise RuntimeError("Authoritative rev5 changed during rev6 construction")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--benchmark-summary", required=True)
    parser.add_argument("--parameter-counts", required=True)
    parser.add_argument("--figure", required=True)
    return parser.parse_args()


if __name__ == "__main__":
    build(parse_args())
