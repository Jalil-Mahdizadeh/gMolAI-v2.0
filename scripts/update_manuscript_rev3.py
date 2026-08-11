#!/usr/bin/env python3
"""Create manuscript revision 3 from rev2 and the frozen no-training audits."""

from __future__ import annotations

import argparse
import json
import os
import tempfile
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.shared import Inches
    from docx.text.paragraph import Paragraph
except ImportError as error:  # pragma: no cover - document build dependency
    raise SystemExit("python-docx>=1.2 is required to build manuscript rev3") from error


def load_json(path: Path) -> dict[str, Any]:
    if not path.is_file():
        raise FileNotFoundError(path)
    return json.loads(path.read_text(encoding="utf-8"))


def find_paragraph(document, prefix: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {prefix!r}, found {len(matches)}")
    return matches[0]


def replace_phrase(document, old: str, new: str) -> None:
    matches = [paragraph for paragraph in document.paragraphs if old in paragraph.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one occurrence of {old!r}, found {len(matches)}")
    paragraph = matches[0]
    run_matches = [run for run in paragraph.runs if old in run.text]
    if len(run_matches) == 1:
        run_matches[0].text = run_matches[0].text.replace(old, new)
    else:
        paragraph.text = paragraph.text.replace(old, new)


def replace_paragraph(document, prefix: str, text: str) -> Paragraph:
    paragraph = find_paragraph(document, prefix)
    paragraph.text = text
    return paragraph


def insert_after(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    inserted.style = style
    inserted.add_run(text)
    return inserted


def metric_summary(control: dict[str, Any], dataset: str) -> tuple[str, dict[str, tuple[float, float]]]:
    item = control["datasets"][dataset]
    metric = item["comparison"]["primary_metric"]
    values = {}
    for feature in ("molecule_embedding", "morgan_radius2_2048", "auxiliary_descriptors_13"):
        summary = item["feature_results"][feature]["summary"][metric]
        values[feature] = (float(summary["mean"]), float(summary["std"]))
    return metric, values


def fmt(mean: float, std: float) -> str:
    return f"{mean:.4f} ± {std:.4f}"


def update_downstream_table(document, control: dict[str, Any]) -> None:
    matches = [
        table
        for table in document.tables
        if table.rows and table.rows[0].cells[0].text.strip() == "Dataset (n)"
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one downstream table, found {len(matches)}")
    table = matches[0]
    if len(table.rows) != 6 or len(table.columns) != 6:
        raise RuntimeError("Unexpected rev2 downstream table shape")

    table.add_column(Inches(0.95))
    for row in table.rows:
        cells = list(row._tr.tc_lst)
        cells[4].addprevious(cells[-1])
    header = [
        "Dataset (n)",
        "Primary metric",
        "gMolAI-v2.0",
        "Morgan r=2, 2048-bit",
        "13 descriptors",
        "Promotion gate",
        "Numerical comparison",
    ]
    for index, value in enumerate(header):
        table.rows[0].cells[index].text = value

    row_order = ("bace", "bbbp", "esol", "freesolv", "lipophilicity")
    expected_labels = {
        "bace": "BACE (1,513)",
        "bbbp": "BBBP (1,860)",
        "esol": "ESOL (1,116)",
        "freesolv": "FreeSolv (639)",
        "lipophilicity": "Lipophilicity (4,198)",
    }
    comparisons = {
        "bace": "Morgan +0.0147; descriptors lower",
        "bbbp": "gMolAI +0.0139; descriptors lower",
        "esol": "gMolAI 0.1457 lower than descriptors",
        "freesolv": "gMolAI 0.3047 lower than descriptors",
        "lipophilicity": "gMolAI 0.1946 lower than descriptors",
    }
    for row_index, dataset in enumerate(row_order, start=1):
        row = table.rows[row_index]
        if row.cells[0].text.strip() != expected_labels[dataset]:
            raise RuntimeError(f"Unexpected downstream table row {row_index}")
        _, values = metric_summary(control, dataset)
        row.cells[2].text = fmt(*values["molecule_embedding"])
        row.cells[3].text = fmt(*values["morgan_radius2_2048"])
        row.cells[4].text = fmt(*values["auxiliary_descriptors_13"])
        row.cells[6].text = comparisons[dataset]

    widths = (1.00, 0.75, 0.90, 0.95, 0.90, 0.75, 1.20)
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = 0


def validate_inputs(control: dict[str, Any], overlap: dict[str, Any], exposure: dict[str, Any]) -> None:
    if control.get("pretrained_model_executed") is not False:
        raise RuntimeError("Descriptor control does not certify model-free execution")
    if int(control["reference_benchmark"]["checkpoint"]["global_step"]) != 10_000:
        raise RuntimeError("Descriptor control is not bound to checkpoint 10k")
    if overlap.get("pretrained_model_executed") is not False:
        raise RuntimeError("Overlap audit does not certify model-free execution")
    if exposure.get("pretrained_model_executed") is not False:
        raise RuntimeError("Exposure audit does not certify model-free execution")
    if exposure.get("training_permitted") is not False:
        raise RuntimeError("Exposure audit does not certify the no-training constraint")
    if int(exposure.get("training_stream_seed", -1)) != 42:
        raise RuntimeError("Exposure audit is not bound to the seed-42 training stream")
    steps = {int(item["global_step"]): item for item in exposure["checkpoints"]}
    expected = {10_000: 57_504_265, 15_000: 86_236_032}
    for step, count in expected.items():
        if int(steps[step]["unique_training_graphs_presented"]) != count:
            raise RuntimeError(f"Unexpected exposure count at step {step}")
    for dataset in ("bace", "bbbp", "esol", "freesolv", "lipophilicity"):
        split = control["datasets"][dataset]["split_reconstruction"]
        if split["status"] != "exact_match_to_reference_seeds_and_counts":
            raise RuntimeError(f"{dataset} split reconstruction was not validated")
    expected_overlap = {
        "bace": (414, 413),
        "bbbp": (1_090, 1_080),
        "esol": (969, 964),
        "freesolv": (526, 524),
        "lipophilicity": (2_513, 2_493),
        "hiv": (27_377, 27_145),
    }
    for dataset, (corpus_count, train_count) in expected_overlap.items():
        observed = overlap["datasets"][dataset]["overlap"]
        if (
            int(observed["pretraining_corpus"]) != corpus_count
            or int(observed["by_pretraining_split"]["train"]) != train_count
        ):
            raise RuntimeError(f"Unexpected overlap counts for {dataset}")
    expected_metrics = {
        "bace": ("0.8413 ± 0.0250", "0.8560 ± 0.0366", "0.6889 ± 0.0563"),
        "bbbp": ("0.8789 ± 0.0432", "0.8651 ± 0.0192", "0.8388 ± 0.0505"),
        "esol": ("0.7314 ± 0.0586", "1.5991 ± 0.2060", "0.8771 ± 0.1237"),
        "freesolv": ("1.2972 ± 0.1929", "2.3810 ± 0.4724", "1.6019 ± 0.1574"),
        "lipophilicity": ("0.8125 ± 0.0223", "0.9297 ± 0.0267", "1.0071 ± 0.0380"),
    }
    for dataset, expected_values in expected_metrics.items():
        _, values = metric_summary(control, dataset)
        observed_values = tuple(
            fmt(*values[feature])
            for feature in ("molecule_embedding", "morgan_radius2_2048", "auxiliary_descriptors_13")
        )
        if observed_values != expected_values:
            raise RuntimeError(f"Unexpected downstream metrics for {dataset}")


def build(args: argparse.Namespace) -> None:
    source = Path(args.input).resolve()
    destination = Path(args.output).resolve()
    if source == destination:
        raise ValueError("rev3 output must differ from the retained rev2 input")
    if not source.is_file():
        raise FileNotFoundError(source)
    control = load_json(Path(args.descriptor_control))
    overlap = load_json(Path(args.overlap))
    exposure = load_json(Path(args.exposure))
    validate_inputs(control, overlap, exposure)
    document = Document(source)

    replace_phrase(document, "complete predefined downstream diagnostic panel", "complete downstream diagnostic panel")
    replace_phrase(document, "pre-specified promotion floor", "frozen promotion floor")
    replace_phrase(document, "predefined performance", "frozen performance")
    replace_phrase(document, "Table 5 | Pre-specified promotion criteria.", "Table 5 | Frozen promotion criteria used in the complete retrospective checkpoint audit.")
    replace_phrase(document, "predefined minimum of 25", "frozen minimum of 25")
    replace_phrase(document, "prespecified fail-closed promotion protocol", "frozen fail-closed promotion protocol")
    replace_phrase(document, "predefined 0.82 promotion threshold", "fixed 0.82 promotion threshold")
    replace_phrase(document, "Every predefined representation and development-dataset gate", "Every frozen representation and development-dataset gate")

    replace_paragraph(
        document,
        "Clustering analysis was restricted",
        "Clustering analysis was restricted to recurring Bemis-Murcko scaffolds represented by at least five validation molecules. Up to the 32 most frequent recurring scaffolds were retained. Input vectors for the learned embedding and matched Morgan fingerprints were L2-normalized before standard Euclidean K-means, yielding a cosine-related input geometry; this was not an implementation of spherical K-means because centroids were neither constrained nor renormalized to unit norm. The number of clusters equaled the number of retained scaffolds. Clustering was repeated with five random seeds; each run used 20 K-means initializations and at most 500 iterations. Adjusted Rand index (ARI), normalized mutual information (NMI), homogeneity, completeness and V-measure quantified agreement with scaffold identity. Promotion required learned-embedding ARI to be at least the matched Morgan ARI and learned NMI to be no more than 0.03 below the matched Morgan NMI.",
    )
    replace_paragraph(
        document,
        "Each development task used ten accepted outer scaffold splits.",
        "Each development task used ten accepted outer scaffold splits. GroupShuffleSplit assigned approximately 80% of scaffold groups to the outer training set and 20% of scaffold groups to the outer test set. Because scaffold groups varied in size, the realized molecule fractions differed among splits. For binary classification, candidate outer splits were accepted only when both classes were represented in both partitions and the required inner cross-validation folds were feasible. Hyperparameters were selected only within the outer training partition. Regression used three-fold GroupKFold; classification used three-fold StratifiedGroupKFold with shuffling and deterministic random seeds. Thus, no molecule or scaffold group from an outer test fold contributed to feature scaling or hyperparameter selection for that fold. Across the ten accepted outer splits, summary standard deviations were computed with NumPy population standard deviation (ddof=0). Because repeated GroupShuffleSplit partitions can overlap, these values are descriptive split-to-split dispersions rather than standard errors from ten statistically independent test replicates.",
    )

    diagnostic = find_paragraph(document, "As diagnostic comparators during development")
    heading = insert_after(
        diagnostic,
        "1.6.4. Descriptor-only downstream control and exact pretraining-corpus overlap",
        "Heading 3",
    )
    method_control = insert_after(
        heading,
        "A model-free downstream control used the same 13 auxiliary descriptor definitions as fixed molecular features. For an exact canonical identity present in the immutable pretraining corpus, the stored d00-d12 values were used; identities absent from that corpus were evaluated by recomputing the same frozen definitions with RDKit 2025.09.3. Stored and recomputed values were compared for every overlapping molecule and produced no failures at the deduplication tolerances (absolute 1e-8, relative 1e-5). The original ten accepted outer seeds and train/test counts were reconstructed and validated against the selected 10,000-step benchmark; outer and inner molecule-set identities were then recorded as SHA-256 digests. Descriptor features used the identical fold-local scaling, Ridge or logistic model, hyperparameter grid, inner-fold assignments and metrics as the gMolAI and Morgan evaluations. This control neither loaded nor executed the pretrained model.",
        "Body Text",
    )
    insert_after(
        method_control,
        "Exact molecule overlap between the immutable pretraining corpus and BACE, BBBP, ESOL, FreeSolv, Lipophilicity and HIV was audited using the same downstream canonicalization and deduplication policy. Downstream canonical-isomeric-SMILES SHA-256 identities were joined to the 256 bucketed pretraining Parquet tables, and matches were counted for the complete corpus and separately for its training, validation and internal-test partitions. This audit measures molecular exposure; the experimental downstream endpoint labels were not pretraining targets.",
        "Body Text",
    )

    replace_paragraph(
        document,
        "Retained training milestones were generated at steps 2,500",
        "Retained training milestones were generated at steps 2,500, 5,000, 7,500, 10,000, 12,500 and 15,000. Semantic representation selection was intentionally separate from the online reconstruction-oriented checkpoint score. Repository history shows that the criteria-bearing validator was committed and remained unchanged before the expanded complete sweep. However, the first commit already contains the criteria, the original 10,000-step selection and a preliminary retained-step screen, so it cannot establish prospective specification before those analyses. Accordingly, the 17 frozen promotion criteria were applied uniformly in a complete retrospective audit of retained primary-seed checkpoints from 5,000 through 15,000 steps; the retained 2,500-step checkpoint was not included. Each evaluated checkpoint was paired with a checkpoint-bound calibrator fitted exclusively on 100,000 pretraining-training molecules and was subjected to the same representation probes, similarity and clustering analyses, and ten-split panel on all five development datasets. A candidate could be promoted only if checkpoint, calibrator, graph manifest, configuration, representation exports and downstream panel were mutually identity-consistent, all 17 criteria were satisfied, and every protocol and artifact-integrity requirement passed.",
    )

    replace_paragraph(
        document,
        "The immutable preprocessing pipeline materialized 223,180,699",
        "The model was trained from a corpus of 223,180,699 deduplicated molecular graphs from the combined ZINC and PubChem inputs, including 221,148,895 graphs in the training partition, 943,038 in validation and 1,088,766 in the locked internal test. Exact serialized DDP cursors show that the four-rank seed-42 training stream had presented 57,504,265 unique training graphs by step 10,000 (26.00% of the training partition) and 86,236,032 by step 15,000 (38.99%). All rank cursors remained in their first deterministic shard cycle, so total presentations equaled unique source graphs at both points. Neither checkpoint completed one pass through the training partition. A presentation denotes one source graph consumed into an optimizer batch; masked or corrupted internal views were not counted separately. Training was deliberately step-based and retained immutable milestones every 2,500 steps, separating continued objective optimization from frozen-representation quality.",
    )
    replace_paragraph(
        document,
        "Complete checkpoint evaluation revealed a key decoupling",
        "Complete checkpoint evaluation revealed a key decoupling between continued pretraining improvement and downstream eligibility. The 17 frozen criteria were applied uniformly in the complete retrospective 5,000-15,000-step primary-seed audit, with a separate training-only calibrator and the full representation and five-dataset development panel at each checkpoint. Step 5,000 passed 15 of 17 criteria, failing effective rank (24.865 versus the minimum of 25) and FreeSolv RMSE (1.39045 versus the maximum of 1.30). Steps 7,500, 12,500 and 15,000 each passed 16 of 17 criteria and failed only FreeSolv, with RMSE values of 1.37227, 1.32289 and 1.30346, respectively. Step 10,000 passed all 17 criteria, including FreeSolv RMSE 1.29716, and was the only checkpoint to pass the complete fail-closed gate (Table 6; Fig. 2). Protocol, identity, source-integrity and artifact-completeness checks passed at every evaluated step. No threshold was modified during this expanded audit, and the retained 2,500-step checkpoint was not included.",
    )
    replace_paragraph(
        document,
        "All checkpoints satisfied the protocol, identity",
        "All checkpoints satisfied the protocol, identity, source-integrity and artifact-completeness checks. “Criteria passed” counts the 17 frozen quality criteria in Table 5. FreeSolv failed at steps 5,000, 7,500, 12,500 and 15,000; step 5,000 additionally failed the effective-rank floor. Values were compared at full precision. The retained 2,500-step checkpoint was not included in the complete downstream sweep.",
    )
    replace_paragraph(
        document,
        "Unsupervised K-means analysis provided a complementary global view.",
        "Unsupervised analysis using standard Euclidean K-means on L2-normalized inputs provided a complementary global view; because the fitted centroids were not renormalized, this was not spherical K-means. On the authoritative validation sample, clustering of the gMolAI embeddings gave adjusted Rand index (ARI) 0.362 and normalized mutual information (NMI) 0.788 with respect to recurring Bemis-Murcko scaffolds, compared with ARI 0.317 and NMI 0.777 for the matched Morgan baseline. A second independently sampled 50,000-molecule validation panel again favored gMolAI numerically (ARI 0.352 versus 0.322; NMI 0.785 versus 0.783). These results indicate that scaffold-related organization can emerge from the learned continuous space without requiring it to become an ECFP replica.",
    )

    downstream_intro = find_paragraph(document, "External transfer was assessed using ten repeated")
    split_counts = {
        name: overlap["datasets"][name]["overlap"]
        for name in ("bace", "bbbp", "esol", "freesolv", "lipophilicity", "hiv")
    }
    overlap_text = (
        "Exact canonical-identity auditing found complete-corpus overlap for 414/1,513 BACE, "
        "1,090/1,860 BBBP, 969/1,116 ESOL, 526/639 FreeSolv, 2,513/4,198 "
        "Lipophilicity and 27,377/37,225 HIV molecules. The corresponding overlaps with "
        "the pretraining training partition were 413, 1,080, 964, 524, 2,493 and 27,145. "
        "These molecules were present only with graph-derived and calculated-descriptor "
        "pretraining targets, not the experimental endpoint labels used here. The finding "
        "therefore does not constitute direct endpoint-label leakage, but it does establish "
        "that these downstream evaluations are not molecule-novel relative to the "
        "unlabeled/auxiliary pretraining corpus."
    )
    if split_counts["bace"]["pretraining_corpus"] != 414:
        raise RuntimeError("Overlap artifact changed after manuscript text was prepared")
    insert_after(downstream_intro, overlap_text, "Normal")

    replace_paragraph(
        document,
        "Performance was endpoint dependent (Table 9).",
        "Performance was endpoint dependent (Table 9). On BACE, Morgan achieved the highest mean ROC-AUC (0.8560), followed by gMolAI (0.8413) and the 13-descriptor control (0.6889). On BBBP, gMolAI was numerically highest (0.8789), followed by Morgan (0.8651) and descriptors (0.8388). For ESOL, RMSE was 0.7314 for gMolAI, 0.8771 for descriptors and 1.5991 for Morgan; for FreeSolv the corresponding values were 1.2972, 1.6019 and 2.3810. Thus, relative to Morgan, descriptors accounted descriptively for 83.2% and 71.9% of the mean gMolAI RMSE reduction on ESOL and FreeSolv, while gMolAI remained lower than descriptors by 0.1457 and 0.3047 RMSE. On Lipophilicity, gMolAI achieved RMSE 0.8125, Morgan 0.9297 and descriptors 1.0071; the descriptor control was worse than Morgan and therefore did not explain the gMolAI gain. These paired-split summaries are descriptive and are not tests of statistical superiority or causal decompositions of the learned representation.",
    )
    replace_phrase(
        document,
        "Table 9 | Frozen downstream transfer of the selected 384-D representation versus Morgan fingerprints.",
        "Table 9 | Frozen downstream transfer of the selected 384-D representation, Morgan fingerprints and the 13-descriptor-only control.",
    )
    replace_paragraph(
        document,
        "Values are mean ± descriptive population s.d.",
        "Values are mean ± descriptive population s.d. (NumPy ddof=0) across ten repeated outer scaffold-group splits. GroupShuffleSplit assigned approximately 80%/20% of scaffold groups, whereas realized molecule fractions varied with group size; repeated partitions can overlap, so the reported s.d. is not a standard error from independent test replicates. The 13-descriptor control used immutable stored values for exact corpus matches and pinned-RDKit recomputation for absent identities; all stored-versus-recomputed comparisons passed the configured tolerances. These five datasets were development/promotion tasks and the results are selection-conditioned. Comparisons use matched frozen features and identical downstream protocols; they are not claims of state-of-the-art performance against separately fine-tuned models.",
    )
    replace_paragraph(
        document,
        "The mixed classification results and strong regression results are important",
        "The three-way comparison does not support universal superiority of either the learned representation or an engineered baseline. Morgan remained strongest on BACE, gMolAI was numerically strongest on BBBP and all three regression endpoints, and the 13 descriptors alone were already a strong physicochemical baseline for ESOL and FreeSolv. The residual gMolAI advantage over descriptors on those endpoints, together with the failure of descriptors to match Morgan on Lipophilicity, shows that regression transfer cannot be attributed solely to direct reuse of the auxiliary descriptor vector. The result is nevertheless consistent with meaningful contribution from descriptor supervision and should not be framed as wholly label-free emergence.",
    )
    replace_paragraph(
        document,
        "Because the pretraining objective includes calculated descriptors",
        "Because the pretraining objective includes molecular weight, topological polar surface area, partial charges, logP-related quantities and nine other calculated targets, the descriptor-only control is essential for interpretation. It shows that a large descriptive share of the ESOL and FreeSolv improvement over Morgan is accessible from those 13 values alone, whereas the remaining gMolAI advantage and the Lipophilicity result require information beyond that fixed descriptor vector. This control does not causally separate masked-graph learning from auxiliary supervision because both shaped the trained encoder; an otherwise identical retraining ablation without descriptor loss would be required for causal attribution and was not performed. The held-out topology, retrieval and clustering analyses provide complementary evidence of structural organization beyond the direct descriptor targets.",
    )
    replace_paragraph(
        document,
        "The present results provide three forms of evidence",
        "The present results provide three forms of evidence for added value beyond fixed baselines. First, the representation retained high linear accessibility of topological quantities absent from the pretraining targets. Second, gMolAI nearest-neighbour geometry was chemically enriched but only moderately correlated with Morgan similarity. Third, matched downstream experiments showed endpoint-specific complementarity: Morgan was stronger on BACE, gMolAI was numerically stronger on BBBP, and gMolAI achieved lower mean RMSE than both Morgan and the 13-descriptor-only control on ESOL, FreeSolv and Lipophilicity. The descriptor control substantially narrows the interpretation for ESOL and FreeSolv, while its poor Lipophilicity result shows that the learned regression phenotype is not a simple passthrough of the auxiliary descriptor vector.",
    )
    replace_paragraph(
        document,
        "The complete checkpoint trajectory reinforces",
        "The complete checkpoint trajectory reinforces a broader methodological lesson for molecular pretraining. The 15,000-step model was not numerically collapsed; on several pretraining and dimensional diagnostics it was better than the 10,000-step model, and it passed 16 of the 17 frozen criteria. Nevertheless, its FreeSolv RMSE of 1.30346 remained 0.00346 above the fixed 1.30 ceiling. The same complete fail-closed panel was applied uniformly to every checkpoint from 5,000 through 15,000 steps, and step 10,000 uniquely qualified. Promoting step 15,000 would therefore require a post hoc relaxation of the frozen rule. This eligibility result is not evidence that step 10,000 is statistically superior to step 15,000 on FreeSolv; it is the consequence of the complete retrospective audit. Repository history proves that the validator was committed and unchanged before the expanded sweep, but cannot establish prospective specification before the original selection or preliminary retained-step screen.",
    )
    replace_paragraph(
        document,
        "Second, the pretraining objective is not purely self-supervised",
        "Second, the pretraining objective is not purely self-supervised because it includes 13 calculated molecular descriptors as auxiliary regression targets. The no-model descriptor control demonstrates that these fixed features account for a substantial descriptive fraction of the ESOL and FreeSolv improvement relative to Morgan, although gMolAI remains better than the descriptor control and the Lipophilicity gain is not reproduced by descriptors alone. This comparison is not a causal ablation of the auxiliary loss. Quantifying causal contribution would require retraining an otherwise identical model without descriptor supervision; no such retraining was performed. Exact identity auditing also shows material molecule overlap between the unlabeled/auxiliary pretraining corpus and every downstream dataset. Because endpoint labels were not used in pretraining, this is not direct label leakage, but it limits claims of molecule-level novelty and is disclosed explicitly.",
    )

    update_downstream_table(document, control)

    unsupported = ("pre-specified", "prespecified", "predefined", "predeclared")
    remaining = [
        (index, paragraph.text)
        for index, paragraph in enumerate(document.paragraphs)
        if any(token.lower() in paragraph.text.lower() for token in unsupported)
    ]
    if remaining:
        raise RuntimeError(f"Unsupported prospective terminology remains: {remaining}")

    destination.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary = tempfile.mkstemp(
        prefix=f".{destination.name}.", suffix=".tmp", dir=destination.parent
    )
    os.close(fd)
    try:
        document.save(temporary)
        os.chmod(temporary, source.stat().st_mode & 0o777)
        os.replace(temporary, destination)
    except BaseException:
        try:
            os.unlink(temporary)
        except FileNotFoundError:
            pass
        raise


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--descriptor-control", required=True)
    parser.add_argument("--overlap", required=True)
    parser.add_argument("--exposure", required=True)
    return parser.parse_args()


if __name__ == "__main__":
    build(parse_args())
