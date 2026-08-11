#!/usr/bin/env python3
"""Create manuscript revision 4 from rev3 and the exact exposure audit."""

from __future__ import annotations

import argparse
import json
import os
import tempfile
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt
    from docx.text.paragraph import Paragraph
except ImportError as error:  # pragma: no cover - document build dependency
    raise SystemExit("python-docx>=1.2 is required to build manuscript rev4") from error


DATASET_ORDER = ("bace", "bbbp", "esol", "freesolv", "lipophilicity", "hiv")
DATASET_LABELS = {
    "bace": "BACE",
    "bbbp": "BBBP",
    "esol": "ESOL",
    "freesolv": "FreeSolv",
    "lipophilicity": "Lipo",
    "hiv": "HIV",
}
STEPS = (5_000, 7_500, 10_000, 12_500, 15_000)


def load_json(path: Path) -> dict[str, Any]:
    if not path.is_file():
        raise FileNotFoundError(path)
    return json.loads(path.read_text(encoding="utf-8"))


def find_paragraph(document, prefix: str) -> Paragraph:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(prefix)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph starting {prefix!r}, found {len(matches)}"
        )
    return matches[0]


def replace_paragraph(document, prefix: str, text: str) -> Paragraph:
    paragraph = find_paragraph(document, prefix)
    paragraph.text = text
    return paragraph


def replace_phrase(document, old: str, new: str) -> None:
    matches = [paragraph for paragraph in document.paragraphs if old in paragraph.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one occurrence of {old!r}, found {len(matches)}")
    paragraph = matches[0]
    run_matches = [run for run in paragraph.runs if old in run.text]
    if len(run_matches) == 1:
        run_matches[0].text = run_matches[0].text.replace(old, new)
    else:
        paragraph.text = paragraph.text.replace(old, new)


def insert_after(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    inserted.style = style
    inserted.add_run(text)
    return inserted


def trajectory_item(
    exposure: dict[str, Any], dataset: str, step: int
) -> dict[str, Any]:
    matches = [
        item
        for item in exposure["datasets"][dataset]["trajectory"]
        if int(item["global_step"]) == step
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one {dataset} exposure record at step {step}")
    return matches[0]


def validate_exposure(exposure: dict[str, Any]) -> None:
    immutable_false = (
        "pretrained_model_executed",
        "training_permitted",
        "checkpoints_modified",
        "embeddings_regenerated",
        "promotion_criteria_modified",
        "promoted_seed42_step10000_artifact_modified",
    )
    for key in immutable_false:
        if exposure.get(key) is not False:
            raise RuntimeError(f"Exposure artifact does not certify {key}=false")
    if int(exposure.get("training_stream_seed", -1)) != 42:
        raise RuntimeError("Exposure artifact is not bound to the seed-42 stream")
    checkpoints = {int(item["global_step"]): item for item in exposure["checkpoints"]}
    expected_presentations = {
        5_000: 28_743_683,
        7_500: 43_109_793,
        10_000: 57_504_265,
        12_500: 71_870_280,
        15_000: 86_236_032,
    }
    if set(checkpoints) != set(STEPS):
        raise RuntimeError(
            "Exposure artifact does not contain exactly the five retained steps"
        )
    for step, expected in expected_presentations.items():
        if int(checkpoints[step]["unique_training_graphs_presented"]) != expected:
            raise RuntimeError(f"Unexpected aggregate graph exposure at step {step}")
    expected_overlap = {
        "bace": (1_513, 414, 413),
        "bbbp": (1_860, 1_090, 1_080),
        "esol": (1_116, 969, 964),
        "freesolv": (639, 526, 524),
        "lipophilicity": (4_198, 2_513, 2_493),
        "hiv": (37_225, 27_377, 27_145),
    }
    if set(exposure["datasets"]) != set(DATASET_ORDER):
        raise RuntimeError(
            "Exposure artifact does not contain exactly the six requested datasets"
        )
    for dataset, expected in expected_overlap.items():
        item = exposure["datasets"][dataset]
        observed = (
            int(item["post_filter_downstream_size"]),
            int(item["corpus_overlap"]["count"]),
            int(item["training_partition_overlap"]["count"]),
        )
        if observed != expected or item.get("monotonic") is not True:
            raise RuntimeError(f"Unexpected overlap/exposure validation for {dataset}")
        if [int(row["global_step"]) for row in item["trajectory"]] != list(STEPS):
            raise RuntimeError(f"Incomplete trajectory for {dataset}")
    required_validation = (
        "all_checkpoint_cursors_cycle_zero",
        "all_ranks_accounted_exactly",
        "checkpoint_cursors_monotonic",
        "dataset_seen_sets_monotonic",
        "no_double_counting",
        "canonical_hash_matching_collision_safe",
        "all_training_overlap_locations_resolved_once",
        "rank_shards_exclusive",
    )
    for key in required_validation:
        if exposure["validation"].get(key) is not True:
            raise RuntimeError(f"Exposure artifact failed validation invariant {key}")
    if exposure["validation"].get("tensor_storage_members_loaded") is not False:
        raise RuntimeError("Exposure audit loaded graph tensor storage")
    if int(exposure["validation"]["training_shards_scanned"]) != 27_136:
        raise RuntimeError("Exposure audit did not scan every training shard")
    if int(exposure["validation"]["training_graph_hashes_scanned"]) != 221_148_895:
        raise RuntimeError("Exposure audit did not scan every training graph identity")


def _format_seen(item: dict[str, Any]) -> str:
    count = int(item["seen_training_overlap_molecules"])
    full = 100.0 * float(item["full_downstream_fraction_seen"])
    train = 100.0 * float(item["training_overlap_fraction_seen"])
    return f"{count:,}\n{full:.2f}/{train:.2f}"


def _set_cell(cell, text: str, *, bold: bool = False, size: float = 6.6) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    properties = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", 25), ("left", 25), ("bottom", 25), ("right", 25)):
        margin = OxmlElement(f"w:{side}")
        margin.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w",
            str(value),
        )
        margin.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "dxa"
        )
        margins.append(margin)
    properties.append(margins)


def insert_exposure_table(
    document, anchor: Paragraph, exposure: dict[str, Any]
) -> None:
    caption_style = find_paragraph(document, "Table 9 |").style.name
    caption = insert_after(
        anchor,
        "Table 10 | Exact checkpoint-resolved downstream-molecule exposure in the seed-42 training stream.",
        caption_style,
    )
    table = document.add_table(rows=1, cols=8)
    table.style = document.tables[8].style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    caption._p.addnext(table._tbl)

    headers = ("Dataset (n)", "Corpus", "Train", "5k", "7.5k", "10k", "12.5k", "15k")
    widths = (0.78, 0.87, 0.87, 0.80, 0.80, 0.80, 0.80, 0.80)
    for index, header in enumerate(headers):
        _set_cell(table.rows[0].cells[index], header, bold=True, size=6.7)
        table.rows[0].cells[index].width = Inches(widths[index])
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "true"
    )
    header_properties.append(repeat)

    display = {
        "bace": "BACE",
        "bbbp": "BBBP",
        "esol": "ESOL",
        "freesolv": "FreeSolv",
        "lipophilicity": "Lipo",
        "hiv": "HIV*",
    }
    for dataset in DATASET_ORDER:
        item = exposure["datasets"][dataset]
        accepted = int(item["post_filter_downstream_size"])
        corpus = item["corpus_overlap"]
        train = item["training_partition_overlap"]
        values = [
            f"{display[dataset]}\n({accepted:,})",
            f"{int(corpus['count']):,}\n{100.0 * float(corpus['fraction_of_downstream']):.2f}",
            f"{int(train['count']):,}\n{100.0 * float(train['fraction_of_downstream']):.2f}",
            *[_format_seen(trajectory_item(exposure, dataset, step)) for step in STEPS],
        ]
        row = table.add_row()
        for index, value in enumerate(values):
            _set_cell(row.cells[index], value, size=6.5)
            row.cells[index].width = Inches(widths[index])

    note = document.add_paragraph(style="Normal")
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(6)
    run = note.add_run(
        "Corpus and Train cells give n and % of the post-filter dataset. Checkpoint cells give "
        "the exact number consumed before the checkpoint and, on the second line, % of the "
        "full downstream dataset/% of training-partition overlaps. *HIV was confirmatory and "
        "not used for checkpoint selection."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(7)
    table._tbl.addnext(note._p)


def _dataset_counts_sentence(exposure: dict[str, Any], step: int) -> str:
    values = []
    for dataset in DATASET_ORDER:
        seen = trajectory_item(exposure, dataset, step)
        values.append(
            f"{DATASET_LABELS[dataset]} {int(seen['seen_training_overlap_molecules']):,} "
            f"({100.0 * float(seen['training_overlap_fraction_seen']):.2f}% of training overlaps)"
        )
    return ", ".join(values[:-1]) + f", and {values[-1]}"


def build(args: argparse.Namespace) -> None:
    source = Path(args.input).resolve()
    destination = Path(args.output).resolve()
    if source == destination:
        raise ValueError("rev4 output must differ from the retained rev3 input")
    if not source.is_file():
        raise FileNotFoundError(source)
    exposure = load_json(Path(args.downstream_exposure))
    validate_exposure(exposure)
    document = Document(source)

    replace_paragraph(
        document,
        "1.6.4. Descriptor-only downstream control and exact pretraining-corpus overlap",
        "1.6.4. Descriptor-only downstream control, corpus overlap and checkpoint-resolved exposure",
    )
    overlap_method = find_paragraph(document, "Exact molecule overlap between")
    insert_after(
        overlap_method,
        "Actual checkpoint-specific exposure was then reconstructed without training, model "
        "execution or embedding regeneration for retained seed-42 steps 5,000, 7,500, 10,000, "
        "12,500 and 15,000. Each checkpoint was verified against the immutable graph-manifest "
        "identity and its four rank-specific cursors. Training shards were assigned exclusively "
        "by manifest index modulo world size; cycle-0 shard permutations and within-shard graph "
        "permutations were reproduced from the recorded seed, rank, cycle and stable shard-path "
        "hash. A restricted metadata-only reader extracted the ordered full SHA-256 molecular "
        "identities from every training shard while skipping tensor-storage members. All 27,136 "
        "training shards and 221,148,895 graph identities were checked, and every downstream "
        "training-overlap identity mapped to exactly one graph boundary. For the current shard, "
        "a molecule was counted only when its shuffled position was strictly less than the saved "
        "graph_position, which identifies the next unread graph. Nested seen-identity sets, "
        "exclusive rank accounting, canonical-SMILES equality after hash matching and absence of "
        "double counting were enforced as fail-closed invariants.",
        "Body Text",
    )

    replace_paragraph(
        document,
        "The model was trained from a corpus of 223,180,699",
        "The model was trained from a corpus of 223,180,699 deduplicated molecular graphs from "
        "the combined ZINC and PubChem inputs, including 221,148,895 graphs in the training "
        "partition, 943,038 in validation and 1,088,766 in the locked internal test. Exact "
        "serialized DDP cursors show that the four-rank seed-42 stream had presented 28,743,683, "
        "43,109,793, 57,504,265, 71,870,280 and 86,236,032 unique training graphs by steps "
        "5,000, 7,500, 10,000, 12,500 and 15,000, respectively, corresponding to 13.00%, 19.49%, "
        "26.00%, 32.50% and 38.99% of the training partition. Every cursor remained in cycle 0, "
        "so total presentations equaled unique source graphs and no retained checkpoint completed "
        "one training-partition pass. A presentation denotes one source graph included in a "
        "completed optimizer batch; masked or corrupted internal views were not counted "
        "separately. Training was deliberately step-based and retained immutable milestones every "
        "2,500 steps, separating continued objective optimization from frozen-representation quality.",
    )

    replace_paragraph(
        document,
        "Exact canonical-identity auditing found complete-corpus overlap",
        "Exact canonical-identity auditing found complete-corpus overlap for 414/1,513 BACE, "
        "1,090/1,860 BBBP, 969/1,116 ESOL, 526/639 FreeSolv, 2,513/4,198 Lipophilicity and "
        "27,377/37,225 HIV molecules; corresponding training-partition overlaps were 413, 1,080, "
        "964, 524, 2,493 and 27,145. Checkpoint-resolved reconstruction refined this static "
        "membership result: at the promoted step 10,000, exact consumed counts were "
        f"{_dataset_counts_sentence(exposure, 10_000)}. By step 15,000 the corresponding counts "
        f"were {_dataset_counts_sentence(exposure, 15_000)}. Every five-point trajectory was "
        "monotonic (Table 10). These molecules carried graph-derived and calculated-descriptor "
        "pretraining targets, not the experimental endpoint labels; the finding is therefore not "
        "direct endpoint-label leakage. It does, however, show that a measured subset of each "
        "evaluation dataset had actually influenced the pretraining updates, while the remainder "
        "was either only present outside the pretraining training partition or had not yet been "
        "consumed by the checkpoint.",
    )

    anchor = find_paragraph(
        document, "Because the pretraining objective includes molecular weight"
    )
    insert_exposure_table(document, anchor, exposure)
    replace_phrase(
        document,
        "Table 10 | Cross-training-seed replication at step 10,000.",
        "Table 11 | Cross-training-seed replication at step 10,000.",
    )
    replace_paragraph(
        document,
        "Second, the pretraining objective is not purely self-supervised",
        "Second, the pretraining objective is not purely self-supervised because it includes 13 "
        "calculated molecular descriptors as auxiliary regression targets. The no-model descriptor "
        "control demonstrates that these fixed features account for a substantial descriptive "
        "fraction of the ESOL and FreeSolv improvement relative to Morgan, although gMolAI remains "
        "better than the descriptor control and the Lipophilicity gain is not reproduced by "
        "descriptors alone. This comparison is not a causal ablation; quantifying causal "
        "contribution would require retraining without descriptor supervision, which was not "
        "performed. Exact identity auditing further distinguishes static corpus membership from "
        "actual optimizer exposure: only the checkpoint-specific subsets reported in Table 10 had "
        "been consumed by each retained model. Endpoint labels were never pretraining targets, so "
        "this is not direct label leakage. Nevertheless, actual exposure of a nonzero subset at the "
        "promoted checkpoint limits claims of molecule-level novelty and should be considered when "
        "interpreting all six downstream results; the audit does not establish that exposure caused "
        "the observed endpoint performance.",
    )

    if len(document.tables) != 11:
        raise RuntimeError(
            f"Expected 11 publication tables in rev4, found {len(document.tables)}"
        )
    if len([p for p in document.paragraphs if p.text.startswith("Table 10 |")]) != 1:
        raise RuntimeError("rev4 Table 10 caption is not unique")
    if len([p for p in document.paragraphs if p.text.startswith("Table 11 |")]) != 1:
        raise RuntimeError("rev4 Table 11 caption is not unique")

    destination.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary = tempfile.mkstemp(
        prefix=f".{destination.name}.", suffix=".tmp", dir=destination.parent
    )
    os.close(fd)
    try:
        document.save(temporary)
        os.chmod(temporary, source.stat().st_mode & 0o777)
        os.replace(temporary, destination)
    except BaseException:
        try:
            os.unlink(temporary)
        except FileNotFoundError:
            pass
        raise


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--downstream-exposure", required=True)
    return parser.parse_args()


if __name__ == "__main__":
    build(parse_args())
